# utils.py - VERSIÓN COMPLETA Y CORREGIDA

import streamlit as st
import re
import os
import io
import json
import docx
import imgkit
from pypdf import PdfReader
import pandas as pd
import time
import google.api_core.exceptions
from PIL import Image
import google.generativeai as genai

# Importación desde tus módulos
from drive_utils import find_or_create_folder, get_or_create_lot_folder_id, clean_folder_name

# =============================================================================
#           FUNCIONES DE PROCESAMIENTO DE TEXTO Y JSON
# =============================================================================

# ¡AQUÍ ESTÁN TUS CONSTANTES CALIBRADAS!
CARACTERES_POR_PAGINA_MIN = 2100
CARACTERES_POR_PAGINA_MAX = 2200

CONTEXTO_LOTE_TEMPLATE = """

**INSTRUCCIÓN CRÍTICA DE ANÁLIS:** Tu análisis debe centrarse única y exclusivamente en la información relacionada con el **'{lote_seleccionado}'**. Ignora por completo cualquier dato, requisito o criterio de valoración que pertenezca a otros lotes.

"""
OPCION_ANALISIS_GENERAL = "Análisis general (no centrarse en un lote)"

def get_lot_index_info(service, project_folder_id, selected_lot):
    """
    Calcula el ID de la carpeta y el nombre de archivo específico para el índice.
    """
    is_general_analysis = (selected_lot == OPCION_ANALISIS_GENERAL or selected_lot is None)

    if is_general_analysis:
        app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=project_folder_id)
        index_filename = "ultimo_indice.json"
        target_folder_id = app_folder_id
    else:
        lot_folder_id = get_or_create_lot_folder_id(service, project_folder_id, lot_name=selected_lot)
        target_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=lot_folder_id)
        match = re.search(r'Lote\s*(\d+|\w+)', selected_lot, re.IGNORECASE)
        lot_suffix = match.group(1).replace(' ', '_') if match else clean_folder_name(selected_lot).split('_')[0]
        index_filename = f"ultimo_indice_lote{lot_suffix}.json"

    return target_folder_id, index_filename

def get_lot_context():
    """Genera el texto de contexto para la IA si hay un lote seleccionado."""
    lote_seleccionado = st.session_state.get('selected_lot')
    if lote_seleccionado and lote_seleccionado != OPCION_ANALISIS_GENERAL:
        return CONTEXTO_LOTE_TEMPLATE.format(lote_seleccionado=lote_seleccionado)
    return ""
    
def enviar_mensaje_con_reintentos(chat, prompt_a_enviar, reintentos=5, delay=60):
    """
    Envía un mensaje a un chat de Gemini, con una lógica de reintentos para errores comunes.
    """
    for i in range(reintentos):
        try:
            response = chat.send_message(prompt_a_enviar)
            return response
        except google.api_core.exceptions.ResourceExhausted as e:
            st.warning(f"⚠️ Límite de la API alcanzado. Reintentando en {delay} segundos... ({i+1}/{reintentos})")
            time.sleep(delay)
        except Exception as e:
            st.error(f"Ocurrió un error inesperado al contactar la API: {e}")
            return None 
    
    st.error("No se pudo obtener una respuesta de la API después de varios intentos.")
    return None
    
def convertir_excel_a_texto_csv(archivo_excel_bytes, nombre_archivo):
    """
    Lee los bytes de un archivo Excel (.xlsx) y los convierte a texto CSV.
    """
    try:
        xls = pd.ExcelFile(archivo_excel_bytes)
        texto_final_csv = ""
        for nombre_hoja in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=nombre_hoja)
            texto_final_csv += f"--- Contenido de la Hoja: '{nombre_hoja}' del archivo '{nombre_archivo}' ---\n"
            texto_final_csv += df.to_csv(index=False)
            texto_final_csv += "\n\n"
        return texto_final_csv
    except Exception as e:
        st.error(f"No se pudo procesar el archivo Excel '{nombre_archivo}': {e}")
        return ""
        
def limpiar_respuesta_json(texto_sucio):
    """
    Extrae un objeto JSON de una cadena de texto potencialmente sucia.
    """
    if not isinstance(texto_sucio, str):
        return ""
    try:
        match = re.search(r'```(json)?\s*(\{.*?\})\s*```', texto_sucio, re.DOTALL)
        if match:
            return match.group(2)
        else:
            start_index = texto_sucio.find('{')
            end_index = texto_sucio.rfind('}')
            if start_index != -1 and end_index != -1 and end_index > start_index:
                return texto_sucio[start_index:end_index + 1]
            return ""
    except Exception:
        return ""

def limpiar_respuesta_final(texto_ia):
    """
    Limpia la respuesta de la IA para la redacción final, eliminando código y frases introductorias.
    """
    if not isinstance(texto_ia, str):
        return ""
    texto_limpio = re.sub(r'```(json|html|mermaid|text)?.*?```', '', texto_ia, flags=re.DOTALL)
    frases_a_eliminar = [
        r'^\s*Aquí tienes el contenido.*?:', r'^\s*Claro, aquí está la redacción para.*?:',
        r'^\s*A continuación se presenta el contenido detallado:', r'^\s*##\s*.*?$'
    ]
    for patron in frases_a_eliminar:
        texto_limpio = re.sub(patron, '', texto_limpio, flags=re.IGNORECASE | re.MULTILINE)
    return texto_limpio.strip()

def corregir_numeracion_markdown(texto_markdown):
    """
    Corrige las listas numeradas en un texto Markdown para que sean consecutivas.
    """
    lineas_corregidas = []
    contador_lista = 0
    en_lista_numerada = False
    for linea in texto_markdown.split('\n'):
        match = re.match(r'^\s*\d+\.\s+', linea)
        if match:
            if not en_lista_numerada:
                en_lista_numerada = True
                contador_lista = 1
            else:
                contador_lista += 1
            texto_del_item = linea[match.end():]
            lineas_corregidas.append(f"{contador_lista}. {texto_del_item}")
        else:
            en_lista_numerada = False
            lineas_corregidas.append(linea)
    return '\n'.join(lineas_corregidas)


# =============================================================================
#           FUNCIONES DE MANIPULACIÓN DE DOCUMENTOS (WORD, HTML)
# =============================================================================

def natural_sort_key(s):
    """Crea una clave para el ordenamiento 'natural' de cadenas."""
    if not isinstance(s, str): return [s]
    return [int(text) if text.isdigit() else text.lower() for text in re.split(r'(\d+)', s)]

def agregar_markdown_a_word(documento, texto_markdown):
    """
    Convierte una cadena de Markdown a un documento de Word, manejando
    encabezados, listas anidadas (con indentación) y negritas.
    """
    patron_encabezado = re.compile(r'^(#+)\s+(.*)')

    def procesar_linea_con_negritas(parrafo, texto):
        # Divide el texto por el marcador de negrita, manteniendo los marcadores
        partes = re.split(r'(\*\*.*?\*\*)', texto)
        for parte in partes:
            if parte.startswith('**') and parte.endswith('**'):
                # Si es negrita, quita los asteriscos y añade el texto en negrita
                parrafo.add_run(parte[2:-2]).bold = True
            elif parte:
                # Si no, añade el texto normal
                parrafo.add_run(parte)

    for linea in texto_markdown.split('\n'):
        linea_limpia = linea.rstrip()
        if not linea_limpia.strip(): continue # Ignorar líneas vacías

        match_encabezado = patron_encabezado.match(linea_limpia.strip())
        if match_encabezado:
            nivel = min(len(match_encabezado.group(1)), 4)
            documento.add_heading(match_encabezado.group(2).strip(), level=nivel)
            continue

        # Detecta el nivel de indentación contando los espacios iniciales
        indentacion = len(linea_limpia) - len(linea_limpia.lstrip(' '))
        nivel_lista = (indentacion // 2) # Asumimos 2 espacios por nivel, puedes ajustarlo a 4

        # Busca patrones de lista (viñeta o numerada)
        match_viñeta = re.match(r'^\s*[\*\-]\s+', linea_limpia)
        match_numerada = re.match(r'^\s*\d+\.\s+', linea_limpia)

        if match_viñeta or match_numerada:
            # Selecciona el estilo de Word basado en el nivel de anidación
            if nivel_lista == 0:
                style = 'List Bullet' if match_viñeta else 'List Number'
            elif nivel_lista == 1:
                style = 'List Bullet 2' if match_viñeta else 'List Number 2'
            else: # Nivel 2 y superiores
                style = 'List Bullet 3' if match_viñeta else 'List Number 3'

            p = documento.add_paragraph(style=style)
            # Limpia el marcador de lista del texto antes de añadirlo
            texto_del_item = re.sub(r'^\s*([\*\-]\s+|\d+\.\s+)', '', linea_limpia)
            procesar_linea_con_negritas(p, texto_del_item)
        else:
            # Si no es una lista, es un párrafo normal
            p = documento.add_paragraph()
            procesar_linea_con_negritas(p, linea_limpia.strip())
            
def generar_indice_word(documento, estructura_memoria):
    """Añade un índice al principio de un documento de Word."""
    documento.add_heading("Índice", level=1)
    if not estructura_memoria:
        documento.add_paragraph("No se encontró una estructura para generar el índice.")
        return
    for seccion in estructura_memoria:
        apartado_titulo = seccion.get("apartado", "Apartado sin título")
        subapartados = seccion.get("subapartados", [])
        p = documento.add_paragraph()
        p.add_run(apartado_titulo).bold = True
        if subapartados:
            for sub in subapartados:
                documento.add_paragraph(f"    {sub}")
    st.toast("Índice generado en el documento.")

# =============================================================================
#           FUNCIONES DE INTERFAZ DE USUARIO (UI)
# =============================================================================

def mostrar_indice_desplegable(estructura, matices=None):
    """Muestra una estructura de índice en Streamlit con apartados desplegables."""
    if not estructura:
        st.warning("La estructura de la memoria está vacía.")
        return

    matices_dict = {item.get('subapartado'): item.get('indicaciones', 'No se encontraron indicaciones.') for item in (matices or [])}

    for seccion in estructura:
        apartado_principal = seccion.get('apartado', 'Sin Título')
        with st.expander(f"**{apartado_principal}**"):
            subapartados = seccion.get('subapartados', [])
            if not subapartados:
                st.write("No hay subapartados.")
            else:
                for subapartado in subapartados:
                    st.markdown(f" L &nbsp; **{subapartado}**")
                    if indicaciones := matices_dict.get(subapartado):
                        with st.container(border=True):
                            st.info(indicaciones)
                    st.write("")

# =============================================================================
#           FUNCIONES DE CONVERSIÓN HTML A IMAGEN (CORREGIDAS)
# =============================================================================

def wrap_html_fragment(html_fragment):
    """
    (VERSIÓN CORREGIDA)
    Envuelve un fragmento de HTML en una estructura completa con estilos CSS.
    Se elimina la dependencia de Google Fonts (@import) para evitar errores de red.
    """
    if html_fragment.strip().startswith('<!DOCTYPE html>'):
        return html_fragment
    
    css_styles = """
        body { 
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif, "Apple Color Emoji", "Segoe UI Emoji"; 
            background-color: #f0f2f5; 
            display: flex; 
            justify-content: center; 
            align-items: center; 
            padding: 20px; 
            width: 800px; 
            box-sizing: border-box; 
        }
        .card { background-color: white; border-radius: 10px; box-shadow: 0 4px 8px rgba(0,0,0,0.1); padding: 25px; width: 100%; max-width: 750px; border-top: 5px solid #0046C6; }
        h2 { color: #0046C6; text-align: center; margin-top: 0; font-size: 24px; font-weight: 700; }
        ul { list-style-type: none; padding: 0; }
        li { display: flex; align-items: center; margin-bottom: 15px; font-size: 16px; color: #333; }
        li::before { content: '✔'; color: #32CFAA; font-size: 20px; font-weight: bold; margin-right: 15px; }
        table { width: 100%; border-collapse: collapse; margin-top: 20px; font-size: 15px; }
        th, td { padding: 12px 15px; border: 1px solid #ddd; text-align: left; }
        th { background-color: #f5f5f5; font-weight: 600; color: #333; }
        tr:nth-child(even) { background-color: #f9f9f9; }
    """
    
    return f"""
    <!DOCTYPE html><html lang="es"><head><meta charset="UTF-8"><title>Visual Element</title><style>{css_styles}</style></head>
    <body>{html_fragment}</body></html>
    """

def html_a_imagen(html_string):
    """
    (VERSIÓN CORREGIDA)
    Convierte una cadena de HTML en bytes de una imagen PNG, operando en memoria.
    Devuelve los bytes de la imagen si tiene éxito, o None si falla.
    """
    try:
        path_wkhtmltoimage = os.popen('which wkhtmltoimage').read().strip()
        if not path_wkhtmltoimage:
            st.error("❌ 'wkhtmltoimage' no encontrado. Asegúrate de que 'wkhtmltopdf' está en tu archivo packages.txt.")
            return None
        config = imgkit.config(wkhtmltoimage=path_wkhtmltoimage)
        options = {'format': 'png', 'encoding': "UTF-8", 'width': '800', 'quiet': ''}
        
        # El cambio clave: output_path=False devuelve los bytes directamente
        image_bytes = imgkit.from_string(html_string, False, config=config, options=options)
        
        return image_bytes
        
    except Exception as e:
        st.error(f"Error al convertir HTML a imagen: {e}")
        return None

# =============================================================================
#         ARQUITECTURA DE ANÁLISIS MULTIMODAL CON CACHÉ
# =============================================================================

def _analizar_docx_core(file_bytes_io, nombre_archivo):
    """(FUNCIÓN INTERNA) Analiza un .docx extrayendo texto e imágenes."""
    try:
        doc = docx.Document(file_bytes_io)
        prompt_parts = [
            "Eres un analista experto de documentos de licitación. Analiza el siguiente contenido de un documento (texto e imágenes) y genera un resumen de texto enriquecido que capture toda la información clave, describiendo las imágenes en su contexto.\n\n--- INICIO DEL DOCUMENTO ---\n"
        ]
        texto_completo = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        prompt_parts.append(texto_completo)
        prompt_parts.append("\n--- FIN DEL TEXTO / INICIO DE LAS IMÁGENES ---")

        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_bytes = rel.target_part.blob
                    img = Image.open(io.BytesIO(image_bytes))
                    img.thumbnail((1024, 1024))
                    prompt_parts.append(img)
                    image_count += 1
                except Exception:
                    pass # Omitir imágenes no soportadas
        
        if not texto_completo and image_count == 0: return ""

        safety_settings = [{"category": c, "threshold": "BLOCK_NONE"} for c in ["HARM_CATEGORY_HARASSMENT", "HARM_CATEGORY_HATE_SPEECH", "HARM_CATEGORY_SEXUALLY_EXPLICIT", "HARM_CATEGORY_DANGEROUS_CONTENT"]]
        
        model = st.session_state.gemini_model
        response = model.generate_content(prompt_parts, safety_settings=safety_settings)

        if not response.candidates:
            reason = response.prompt_feedback.block_reason.name if hasattr(response, 'prompt_feedback') else "No especificado"
            return f"Error: La API bloqueó la respuesta. Razón: {reason}"

        return f"--- ANÁLISIS MULTIMODAL DE '{nombre_archivo}' ---\n{response.text}"

    except Exception as e:
        return f"Error al procesar el archivo DOCX: {str(e)}"

@st.cache_data(show_spinner=False)
def get_cached_multimodal_analysis(_file_content_bytes, nombre_archivo):
    """(FUNCIÓN CACHEABLE) Envuelve la lógica de análisis principal."""
    print(f"CACHE MISS: Ejecutando análisis por primera vez para '{nombre_archivo}'.")
    return _analizar_docx_core(io.BytesIO(_file_content_bytes), nombre_archivo)

def analizar_docx_multimodal_con_gemini(file_bytes_io, nombre_archivo):
    """(FUNCIÓN PRINCIPAL) Llama a la función cacheable para obtener el análisis."""
    with st.spinner(f"Analizando '{nombre_archivo}' (texto e imágenes)..."):
        analysis_result = get_cached_multimodal_analysis(file_bytes_io.getvalue(), nombre_archivo)
        if "Error" in analysis_result:
            st.error(f"No se pudo analizar '{nombre_archivo}'.")
            return None
        st.success(f"Análisis de '{nombre_archivo}' completado.")
        return analysis_result

def generar_fragmento_individual(model, prompt_info, reintentos=5, delay_inicial=60):
    """
    (VERSIÓN MEJORADA CON REINTENTOS)
    Función segura para hilos que genera un único fragmento de texto.
    Implementa una estrategia de reintentos con exponential backoff para errores de API.
    """
    subapartado = prompt_info.get("subapartado_referencia", "Desconocido")
    prompt_a_enviar = prompt_info.get("prompt_para_asistente")
    prompt_id = prompt_info.get("prompt_id")

    if not prompt_a_enviar:
        return {'success': False, 'error': 'El prompt estaba vacío.', 'prompt_id': prompt_id}

    for i in range(reintentos):
        try:
            # Cada llamada es independiente
            response = model.generate_content(prompt_a_enviar)
            
            if not response.candidates:
                reason = "Bloqueado por filtros de seguridad"
                if hasattr(response, 'prompt_feedback') and response.prompt_feedback.block_reason:
                    reason = response.prompt_feedback.block_reason.name
                # Este no es un error de API, no reintentamos
                return {'success': False, 'error': f"Respuesta bloqueada ({reason})", 'prompt_id': prompt_id}
            
            # Éxito, devolvemos el resultado
            return {'success': True, 'content': response.text, 'prompt_id': prompt_id}

        except google.api_core.exceptions.ResourceExhausted as e:
            delay = delay_inicial * (2 ** i)  # Exponential backoff
            print(f"HILO {prompt_id}: Límite de API alcanzado. Reintentando en {delay} segundos... (Intento {i+1}/{reintentos})")
            time.sleep(delay)

        except Exception as e:
            # Para otros errores, fallamos directamente sin reintentar
            print(f"HILO {prompt_id}: Error inesperado. {str(e)}")
            return {'success': False, 'error': str(e), 'prompt_id': prompt_id}
    
    # Si todos los reintentos fallan
    return {'success': False, 'error': f"Límite de API excedido tras {reintentos} intentos.", 'prompt_id': prompt_id}


# -----------------------------------------------------------------------------
#      ¡AQUÍ ESTÁ LA FUNCIÓN QUE FALTABA!
# -----------------------------------------------------------------------------
def apply_safety_margin_to_plan(generated_structure, safety_margin_factor=0.85):
    """
    Aplica un margen de seguridad al plan de extensión de un índice generado.
    Recalcula las páginas y los caracteres de forma proporcional.

    Args:
        generated_structure (dict): El objeto JSON completo generado por la IA.
        safety_margin_factor (float): El factor por el cual reducir el objetivo (ej. 0.85 para un 15% de reducción).

    Returns:
        dict: El objeto JSON con el plan_extension ajustado.
    """
    try:
        # Extraemos el plan y la configuración originales
        original_plan = generated_structure.get("plan_extension", [])
        config = generated_structure.get("configuracion_licitacion", {})
        
        # Intentamos obtener el máximo de páginas como un número
        max_paginas_str = config.get('max_paginas', 'N/D')
        # Usamos regex para encontrar el número en el string (ej. "20 páginas")
        match = re.search(r'\d+', str(max_paginas_str))
        if not match:
            # Si no hay un número de páginas, no podemos ajustar nada. Devolvemos el original.
            return generated_structure

        original_total_pages = int(match.group(0))
        
        # Si el plan ya es muy corto, no lo reducimos más
        if original_total_pages <= 3:
            return generated_structure

        # Calculamos el nuevo objetivo de páginas
        adjusted_total_pages = original_total_pages * safety_margin_factor

        # Calculamos el total de páginas sugeridas en el plan original de la IA
        # para mantener la proporcionalidad. A veces la IA no suma exactamente el máximo.
        current_plan_total_pages = sum(item.get('paginas_sugeridas_apartado', 0) for item in original_plan)
        if current_plan_total_pages == 0:
            return generated_structure # Evitar división por cero

        # Creamos el nuevo plan que reemplazará al original
        new_plan = []
        
        # Recorremos cada APARTADO del plan
        for apartado_item in original_plan:
            original_apartado_pages = apartado_item.get('paginas_sugeridas_apartado', 0)
            
            # Calculamos las nuevas páginas para este apartado de forma proporcional
            scaling_factor = original_apartado_pages / current_plan_total_pages
            adjusted_apartado_pages = adjusted_total_pages * scaling_factor
            
            new_apartado_item = apartado_item.copy()
            new_apartado_item['paginas_sugeridas_apartado'] = round(adjusted_apartado_pages, 2)
            
            new_subapartado_desglose = []
            original_subapartado_total_pages = sum(sub.get('paginas_sugeridas', 0) for sub in apartado_item.get('desglose_subapartados', []))
            if original_subapartado_total_pages == 0:
                continue

            # Recorremos cada SUBAPARTADO dentro del apartado
            for subapartado_item in apartado_item.get('desglose_subapartados', []):
                original_sub_pages = subapartado_item.get('paginas_sugeridas', 0)
                
                # Calculamos las nuevas páginas para este subapartado proporcionalmente
                sub_scaling_factor = original_sub_pages / original_subapartado_total_pages
                adjusted_sub_pages = adjusted_apartado_pages * sub_scaling_factor
                
                new_sub_item = subapartado_item.copy()
                new_sub_item['paginas_sugeridas'] = round(adjusted_sub_pages, 2)
                
                # ¡Recalculamos los caracteres con los valores de tu propio código!
                new_sub_item['min_caracteres_sugeridos'] = int(adjusted_sub_pages * CARACTERES_POR_PAGINA_MIN)
                new_sub_item['max_caracteres_sugeridos'] = int(adjusted_sub_pages * CARACTERES_POR_PAGINA_MAX)
                
                new_subapartado_desglose.append(new_sub_item)
            
            new_apartado_item['desglose_subapartados'] = new_subapartado_desglose
            new_plan.append(new_apartado_item)

        # Reemplazamos el plan antiguo por el nuevo en la estructura final
        generated_structure["plan_extension"] = new_plan
        st.toast(f"✅ Margen de seguridad aplicado. Objetivo ajustado de {original_total_pages} a ~{int(adjusted_total_pages)} páginas.")
        
        return generated_structure

    except Exception as e:
        st.warning(f"No se pudo aplicar el margen de seguridad: {e}")
        # Si algo falla, devolvemos la estructura original para no romper la app
        return generated_structure


def desensamblar_docx(docx_buffer):
    """
    Analiza un documento de Word desde un buffer, extrae todo el texto y mapea las imágenes.
    Devuelve el texto con placeholders y un diccionario que asocia cada placeholder con su imagen.
    """
    texto_para_ia = ""
    mapa_de_imagenes = {}
    image_counter = 0
    
    docx_buffer.seek(0)
    doc = docx.Document(docx_buffer)

    # El cuerpo del documento contiene párrafos y tablas.
    # Esta es una forma simplificada de recorrerlo. Una implementación real puede necesitar ser más robusta.
    for p in doc.paragraphs:
        # Busca imágenes dentro de cada párrafo
        if 'graphicData' in p._p.xml:
            # Encuentra todos los IDs de relación de imagen en el XML del párrafo
            rids = p._p.xpath('.//a:blip/@r:embed')
            for rid in rids:
                image_counter += 1
                placeholder = f"[--IMAGEN_ID_{image_counter:03d}--]"
                
                # Obtiene la imagen real usando el ID de relación
                image_part = doc.part.related_parts[rid]
                image_blob = image_part.blob
                
                mapa_de_imagenes[placeholder] = image_blob
                texto_para_ia += f"\n{placeholder}\n"
        else:
            # Si no hay imagen, simplemente añade el texto del párrafo
            texto_para_ia += p.text + "\n"
            
    return texto_para_ia, mapa_de_imagenes


def reensamblar_docx_con_imagenes(texto_cohesionado, mapa_de_imagenes):
    """
    Toma el texto procesado por la IA (con placeholders) y el mapa de imágenes,
    y reconstruye un nuevo documento de Word con el texto y las imágenes en su lugar.
    """
    doc_final = docx.Document()
    
    # Divide el texto por el patrón de los placeholders, manteniendo los placeholders en la lista
    fragmentos = re.split(r'(\[--IMAGEN_ID_\d{3}--\])', texto_cohesionado)
    
    for fragmento in fragmentos:
        if not fragmento:
            continue
            
        if fragmento in mapa_de_imagenes:
            # Si el fragmento es un placeholder, busca la imagen en el mapa y la añade
            image_data = mapa_de_imagenes[fragmento]
            try:
                # Añade la imagen, estableciendo un ancho estándar para consistencia
                doc_final.add_picture(io.BytesIO(image_data), width=docx.shared.Inches(6.0))
            except Exception as e:
                p_error = doc_final.add_paragraph()
                p_error.add_run(f"[Error al insertar imagen {fragmento}: {e}]").italic = True
        else:
            # Si es texto, utiliza tu función existente para añadirlo con formato Markdown
            agregar_markdown_a_word(doc_final, fragmento)
            
    return doc_final
