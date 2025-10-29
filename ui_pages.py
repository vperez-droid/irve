import streamlit as st
import pandas as pd
import json
import io
import re
import os
import time
import concurrent.futures
import docx
import google.generativeai as genai

# Imports desde tus módulos
from auth import get_credentials

from prompts import (
    PROMPT_DETECTAR_LOTES, PROMPT_REGENERACION, PROMPT_GEMINI_PROPUESTA_ESTRATEGICA, 
    PROMPT_CONSULTOR_REVISION, PROMPT_DESARROLLO, PROMPT_CLASIFICAR_DOCUMENTO, PROMPT_GENERAR_INTRODUCCION, 
    PROMPT_PLIEGOS, PROMPT_REQUISITOS_CLAVE
)
from drive_utils import (
    find_or_create_folder, get_files_in_project, delete_file_from_drive,
    upload_file_to_drive, find_file_by_name, download_file_from_drive_cached, download_file_from_drive_uncached,
    sync_guiones_folders_with_index, list_project_folders, ROOT_FOLDER_NAME,
    get_or_create_lot_folder_id, clean_folder_name, get_context_from_lots
)
from utils import (
    mostrar_indice_desplegable, limpiar_respuesta_json, agregar_markdown_a_word,
    wrap_html_fragment, html_a_imagen, limpiar_respuesta_final, analizar_docx_multimodal_con_gemini,
    corregir_numeracion_markdown, enviar_mensaje_con_reintentos, get_lot_index_info, generar_indice_word, 
    get_lot_context, OPCION_ANALISIS_GENERAL, natural_sort_key, 
    convertir_excel_a_texto_csv
)


# =============================================================================
#           PÁGINA DE BIENVENIDA / INICIO DE SESIÓN
# =============================================================================

def landing_page():
    from auth import get_google_flow
    col1, col_center, col3 = st.columns([1, 2, 1])
    with col_center:
        st.write("")
        st.markdown(f'<div style="text-align: center;"><img src="https://raw.githubusercontent.com/soporte2-tech/appfront/main/imagen.png" width="150"></div>', unsafe_allow_html=True)
        st.write("")
        st.markdown("<div style='text-align: center;'><h1>Asistente Inteligente para Memorias Técnicas</h1></div>", unsafe_allow_html=True)
        st.markdown("<div style='text-align: center;'><h3>Optimiza y acelera la creación de tus propuestas de licitación</h3></div>", unsafe_allow_html=True)
        st.markdown("---")
        st.info("Para empezar, necesitas dar permiso a la aplicación para que gestione los proyectos en tu Google Drive.")
        flow = get_google_flow()
        auth_url, _ = flow.authorization_url(prompt='consent')
        st.link_button("🔗 Conectar con Google Drive", auth_url, use_container_width=True, type="primary")

# =============================================================================
#           PÁGINA DE SELECCIÓN DE PROYECTO
# =============================================================================

def project_selection_page(go_to_landing, go_to_phase1):
    st.markdown("<h3>Selección de Proyecto</h3>", unsafe_allow_html=True)
    st.markdown("Elige un proyecto existente de tu Google Drive o crea uno nuevo para empezar.")
    st.markdown("---")
    
    service = st.session_state.drive_service
    if not service:
        st.error("No se pudo conectar con Google Drive. Por favor, intenta volver a la página de inicio y reconectar.")
        if st.button("← Volver a Inicio"):
            for key in ['credentials', 'drive_service']:
                if key in st.session_state: del st.session_state[key]
            go_to_landing(); st.rerun()
        return

    with st.spinner("Accediendo a tu Google Drive..."):
        root_folder_id = find_or_create_folder(service, ROOT_FOLDER_NAME)
        projects = list_project_folders(service, root_folder_id)
    
    with st.container(border=True):
        st.subheader("1. Elige un proyecto existente")
        if not projects:
            st.info("Aún no tienes proyectos. Crea uno nuevo en el paso 2.")
        else:
            project_names = ["-- Selecciona un proyecto --"] + list(projects.keys())
            selected_name = st.selectbox("Selecciona tu proyecto:", project_names, key="project_selector")
            
            if st.button("Cargar Proyecto Seleccionado", type="primary"):
                if selected_name != "-- Selecciona un proyecto --":
                    st.session_state.selected_project = {"name": selected_name, "id": projects[selected_name]}
                    st.toast(f"Proyecto '{selected_name}' cargado."); go_to_phase1(); st.rerun()
                else:
                    st.warning("Por favor, selecciona un proyecto de la lista.")

    with st.container(border=True):
        st.subheader("2. O crea un nuevo proyecto")
        new_project_name = st.text_input("Nombre del nuevo proyecto (ej: Licitación Metro Madrid 2024)", key="new_project_name_input")
        if st.button("Crear y Empezar Nuevo Proyecto"):
            if not new_project_name.strip():
                st.warning("Por favor, introduce un nombre para el proyecto.")
            elif new_project_name in projects:
                st.error("Ya existe un proyecto con ese nombre. Por favor, elige otro.")
            else:
                with st.spinner(f"Creando carpeta '{new_project_name}' en tu Drive..."):
                    new_project_id = find_or_create_folder(service, new_project_name, parent_id=root_folder_id)
                    st.session_state.selected_project = {"name": new_project_name, "id": new_project_id}
                    st.success(f"¡Proyecto '{new_project_name}' creado! Ya puedes cargar los documentos.")
                    go_to_phase1(); st.rerun()

# =============================================================================
#           FUNCIÓN phase_1_viability_page
# =============================================================================

def phase_1_viability_page(model, go_to_project_selection, go_to_phase2):
    st.markdown(f"<h3>FASE 1: Análisis de Lotes y Viabilidad</h3>", unsafe_allow_html=True)
    ANALYSIS_FILENAME = "Analisis_de_Viabilidad.docx"
    LOTES_FILENAME = "resultado_analisis_lotes.json"

    if not st.session_state.get('selected_project'):
        st.warning("No se ha seleccionado ningún proyecto. Volviendo a la selección.")
        go_to_project_selection(); st.rerun()
        return
    
    project_name = st.session_state.selected_project['name']
    project_folder_id = st.session_state.selected_project['id']
    service = st.session_state.drive_service
    st.info(f"Proyecto activo: **{project_name}**.")

    with st.container(border=True):
        st.subheader("1. Documentos en tu Proyecto")
        with st.spinner("Buscando archivos en Google Drive..."):
            pliegos_folder_id = find_or_create_folder(service, "Pliegos", parent_id=project_folder_id)
            documentos_pliegos = get_files_in_project(service, pliegos_folder_id)
        if documentos_pliegos:
            st.success("Se analizarán los siguientes archivos encontrados en la carpeta 'Pliegos':")
            for file_info in documentos_pliegos: st.write(f"📄 **{file_info['name']}**")
        else:
            st.warning("No se encontraron archivos en la carpeta 'Pliegos'. Sube al menos un documento para continuar.")
        with st.expander("Subir nuevos documentos a 'Pliegos'"):
            uploaded_files = st.file_uploader("Arrastra aquí los archivos que quieras añadir al proyecto", type=['pdf', 'docx', 'xlsx'], accept_multiple_files=True, key="drive_file_uploader")
            if st.button("Guardar en Drive y Refrescar"):
                if uploaded_files:
                    with st.spinner("Subiendo archivos a Drive..."):
                        for file_obj in uploaded_files: upload_file_to_drive(service, file_obj, pliegos_folder_id)
                    st.toast("¡Archivos subidos!"); st.rerun()

    st.markdown("---")
    
    def detectar_lotes():
        with st.spinner("Analizando documentos para detectar lotes..."):
            try:
                contenido_ia = [PROMPT_DETECTAR_LOTES]
                for file_info in documentos_pliegos:
                    file_bytes_io = download_file_from_drive_cached(service, file_info['id'])
                    nombre_archivo = file_info['name']
                    if nombre_archivo.lower().endswith('.xlsx'):
                        texto_csv = convertir_excel_a_texto_csv(file_bytes_io, nombre_archivo)
                        if texto_csv: contenido_ia.append(texto_csv)
                    else:
                        contenido_ia.append({"mime_type": file_info['mimeType'], "data": file_bytes_io.getvalue()})
                
                response = model.generate_content(contenido_ia, generation_config={"response_mime_type": "application/json"})
                json_limpio = limpiar_respuesta_json(response.text)
                resultado = json.loads(json_limpio)
                lotes = resultado.get("lotes_encontrados", [])
                st.session_state.detected_lotes = lotes if lotes else ["SIN_LOTES"]

                try:
                    docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=project_folder_id)
                    json_bytes = json.dumps(resultado, indent=2).encode('utf-8')
                    mock_file = io.BytesIO(json_bytes)
                    mock_file.name = LOTES_FILENAME
                    mock_file.type = "application/json"
                    
                    existing_file_id = find_file_by_name(service, LOTES_FILENAME, docs_app_folder_id)
                    if existing_file_id:
                        delete_file_from_drive(service, existing_file_id)
                    
                    upload_file_to_drive(service, mock_file, docs_app_folder_id)
                    st.toast("Resultado del análisis de lotes guardado en Drive.")
                except Exception as e:
                    st.warning(f"No se pudo guardar el resultado del análisis de lotes en Drive: {e}")

                st.rerun()
            except Exception as e:
                st.error(f"Ocurrió un error al detectar lotes: {e}")

    st.header("2. Selección de Lote")
    
    if 'detected_lotes' not in st.session_state:
        st.session_state.detected_lotes = None

    if st.session_state.detected_lotes is None:
        try:
            docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=project_folder_id)
            lotes_file_id = find_file_by_name(service, LOTES_FILENAME, docs_app_folder_id)
            
            if lotes_file_id:
                with st.spinner("Cargando análisis de lotes guardado desde Drive..."):
                    file_bytes = download_file_from_drive_cached(service, lotes_file_id).getvalue()
                    resultado = json.loads(file_bytes.decode('utf-8'))
                    lotes = resultado.get("lotes_encontrados", [])
                    st.session_state.detected_lotes = lotes if lotes else ["SIN_LOTES"]
                    st.toast("Análisis de lotes cargado desde Drive.")
                    st.rerun()
        except Exception as e:
            st.warning(f"No se pudo cargar el análisis de lotes guardado. Puede que necesites generarlo de nuevo. Error: {e}")
            st.session_state.detected_lotes = "ERROR"

    if st.session_state.detected_lotes is None or st.session_state.detected_lotes == "ERROR":
        st.info("Antes de analizar la viabilidad, la aplicación comprobará si la licitación está dividida en lotes.")
        st.button("Analizar Lotes en los Documentos", on_click=detectar_lotes, type="primary", use_container_width=True, disabled=not documentos_pliegos)
    
    elif st.session_state.detected_lotes == ["SIN_LOTES"]:
        st.success("✔️ No se han detectado lotes en la documentación. Se realizará un análisis general.")
        if st.session_state.selected_lot is None:
            st.session_state.selected_lot = OPCION_ANALISIS_GENERAL
        st.button("🔄 Forzar Re-análisis de Lotes", on_click=detectar_lotes, help="Vuelve a analizar los documentos si has subido nuevos archivos.", use_container_width=True)

    else:
        st.success("¡Se han detectado lotes en la documentación!")
        if st.session_state.get('selected_lot') is None and st.session_state.detected_lotes:
            st.session_state.selected_lot = st.session_state.detected_lotes[0]
        
        opciones_lotes = st.session_state.detected_lotes + [OPCION_ANALISIS_GENERAL]
        current_selection = st.session_state.get('selected_lot')
        
        try:
            index = opciones_lotes.index(current_selection) if current_selection in opciones_lotes else 0
        except ValueError:
            index = 0
            
        def on_lot_change():
            new_lot = st.session_state.lot_selector_key
            if st.session_state.get('selected_lot') != new_lot:
                st.session_state.selected_lot = new_lot
                if 'analysis_doc_id' in st.session_state: del st.session_state['analysis_doc_id']
                st.toast(f"Lote cambiado a: {new_lot}")

        st.selectbox("Elige el lote al que quieres presentarte o cámbialo si es necesario:", options=opciones_lotes, index=index, key="lot_selector_key", on_change=on_lot_change)
        st.button("🔄 Forzar Re-análisis de Lotes", on_click=detectar_lotes, help="Vuelve a analizar los documentos si has subido nuevos archivos.", use_container_width=True)

    if st.session_state.get('selected_lot') is not None:
        st.markdown("---")
        st.header("3. Extracción de Requisitos Clave")
        
        if st.session_state.selected_lot != OPCION_ANALISIS_GENERAL:
             st.info(f"Se generará el análisis de viabilidad centrado en: **{st.session_state.selected_lot}**")
        else:
             st.info("Se generará un análisis de viabilidad general.")

        selected_lot_name = st.session_state.get('selected_lot')
        active_lot_folder_id = get_or_create_lot_folder_id(service, project_folder_id, lot_name=selected_lot_name)
        docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=active_lot_folder_id)
        
        if 'analysis_doc_id' not in st.session_state:
            st.session_state.analysis_doc_id = find_file_by_name(service, ANALYSIS_FILENAME, docs_app_folder_id)

        def generate_and_save_analysis():
            with st.spinner("🧠 Descargando y analizando documentos con Gemini..."):
                try:
                    idioma = st.session_state.get('project_language', 'Español')
                    contexto_lote = get_lot_context()
                    prompt = PROMPT_REQUISITOS_CLAVE.format(idioma=idioma, contexto_lote=contexto_lote)
                    contenido_ia = [prompt]
                    for file_info in documentos_pliegos:
                        file_bytes_io = download_file_from_drive_cached(service, file_info['id'])
                        nombre_archivo = file_info['name']
                        if nombre_archivo.lower().endswith('.xlsx'):
                            texto_csv = convertir_excel_a_texto_csv(file_bytes_io, nombre_archivo)
                            if texto_csv: contenido_ia.append(texto_csv)
                        else:
                            contenido_ia.append({"mime_type": file_info['mimeType'], "data": file_bytes_io.getvalue()})
                    response = model.generate_content(contenido_ia)
                    if not response.candidates: st.error("Gemini no generó una respuesta."); return
                    documento = docx.Document()
                    agregar_markdown_a_word(documento, response.text)
                    buffer = io.BytesIO()
                    documento.save(buffer); buffer.seek(0)
                    buffer.name = ANALYSIS_FILENAME
                    buffer.type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    if st.session_state.get('analysis_doc_id'):
                        delete_file_from_drive(service, st.session_state['analysis_doc_id'])
                    new_file_id = upload_file_to_drive(service, buffer, docs_app_folder_id)
                    st.session_state.analysis_doc_id = new_file_id
                    st.toast("✅ ¡Análisis guardado en tu Drive!")
                    st.rerun()
                except Exception as e:
                    st.error(f"Ocurrió un error crítico durante el análisis: {e}")

        if st.session_state.analysis_doc_id:
            st.success("✔️ Ya existe un análisis de viabilidad guardado para el lote seleccionado.")
            if st.button("📄 Descargar Análisis Guardado", use_container_width=True):
                with st.spinner("Descargando desde Drive..."):
                    file_bytes = download_file_from_drive_cached(service, st.session_state.analysis_doc_id)
                    st.download_button(label="¡Listo! Haz clic aquí para descargar", data=file_bytes, file_name=ANALYSIS_FILENAME, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            col1, col2 = st.columns(2)
            with col1:
                st.button("🔁 Re-generar Análisis para este Lote", on_click=generate_and_save_analysis, use_container_width=True, disabled=not documentos_pliegos)
            with col2:
                st.button("Continuar a Generación de Índice (Fase 2) →", on_click=go_to_phase2, use_container_width=True, type="primary")
        else:
            st.info("Aún no se ha generado el documento de análisis para este proyecto.")
            st.button("Analizar y Generar Documento de Viabilidad", on_click=generate_and_save_analysis, type="primary", use_container_width=True, disabled=not documentos_pliegos)

    st.write("")
    st.markdown("---")
    st.button("← Volver a Selección de Proyecto", on_click=go_to_project_selection, use_container_width=True)
    
# =============================================================================
#           FASE 2: ANÁLISIS Y ESTRUCTURA
# =============================================================================

def phase_2_structure_page(model, go_to_phase1, go_to_phase2_results, handle_full_regeneration, back_to_project_selection_and_cleanup):
    st.markdown(f"<h3>FASE 2: Análisis y Estructura del Índice</h3>", unsafe_allow_html=True)

    if not st.session_state.get('selected_project'):
        st.warning("No se ha seleccionado ningún proyecto. Volviendo a la selección.")
        return

    project_name = st.session_state.selected_project['name']
    project_folder_id = st.session_state.selected_project['id']
    service = st.session_state.drive_service
    st.info(f"Estás trabajando en el proyecto: **{project_name}**")

    selected_lot = st.session_state.get('selected_lot')
    if selected_lot:
        if selected_lot == OPCION_ANALISIS_GENERAL:
            st.success("🎯 **Enfoque actual:** Se realizará un análisis general para todo el proyecto.")
        else:
            st.success(f"🎯 **Enfoque actual:** Lote / Bloque: **{selected_lot}**")
    else:
        st.warning("No se ha seleccionado un lote o enfoque. Por favor, vuelve a la Fase 1 para continuar.")
        if st.button("← Volver a Fase 1 (Viabilidad)"):
            go_to_phase1(); st.rerun()
        st.stop() 

    st.markdown("---")
    pliegos_folder_id = find_or_create_folder(service, "Pliegos", parent_id=project_folder_id)
    document_files = get_files_in_project(service, pliegos_folder_id)
    
    if document_files:
        st.success("Se usarán estos archivos de la carpeta 'Pliegos' para generar el índice:")
        with st.container(border=True):
            for file in document_files:
                cols = st.columns([4, 1])
                cols[0].write(f"📄 **{file['name']}**")
    else:
        st.info("La carpeta 'Pliegos' de este proyecto está vacía. Sube los archivos base.")

    with st.expander("Añadir o reemplazar documentación en la carpeta 'Pliegos'", expanded=not document_files):
        with st.container(border=True):
            st.subheader("Subir nuevos documentos")
            new_files_uploader = st.file_uploader("Arrastra aquí los nuevos Pliegos o Plantilla", type=['docx', 'pdf', 'xlsx'], accept_multiple_files=True, key="new_files_uploader")
            if st.button("Guardar nuevos archivos en Drive"):
                if new_files_uploader:
                    with st.spinner("Subiendo archivos a la carpeta 'Pliegos'..."):
                        for file_obj in new_files_uploader:
                            upload_file_to_drive(service, file_obj, pliegos_folder_id)
                        st.rerun()
                else:
                    st.warning("Por favor, selecciona al menos un archivo para subir.")

    st.markdown("---"); st.header("Análisis y Generación de Índice")
    
    index_folder_id, index_filename = get_lot_index_info(service, project_folder_id, selected_lot) 
    saved_index_id = find_file_by_name(service, index_filename, index_folder_id)

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Cargar último índice generado", use_container_width=True, disabled=not saved_index_id):
            with st.spinner("Cargando índice desde Drive..."):
                index_content_bytes = download_file_from_drive_cached(service, saved_index_id)
                index_data = json.loads(index_content_bytes.getvalue().decode('utf-8'))
                st.session_state.generated_structure = index_data
                st.session_state.uploaded_pliegos = document_files
                go_to_phase2_results(); st.rerun()

    with col2:
        if st.button("Analizar Archivos y Generar Nuevo Índice", type="primary", use_container_width=True, disabled=not document_files):
            if handle_full_regeneration(model):
                go_to_phase2_results(); st.rerun()

    st.write(""); st.markdown("---")
    
    col_nav1, col_nav2 = st.columns(2)
    with col_nav1:
        st.button("← Volver a Análisis de Viabilidad (F1)", on_click=go_to_phase1, use_container_width=True)
    with col_nav2:
        st.button("↩️ Volver a Selección de Proyecto", on_click=back_to_project_selection_and_cleanup, use_container_width=True, key="back_to_projects")

# =============================================================================
#           FASE 2: REVISIÓN DE RESULTADOS
# =============================================================================

def phase_2_results_page(model, go_to_phase2, go_to_phase3, handle_full_regeneration):
    st.markdown("<h3>FASE 2: Revisión de Resultados del Índice</h3>", unsafe_allow_html=True)
    st.markdown("Revisa el índice, la guía de redacción y el plan estratégico. Puedes hacer ajustes con feedback, regenerarlo todo desde cero, o aceptarlo para continuar.")
    st.markdown("---")
    
    service = st.session_state.drive_service
    project_folder_id = st.session_state.selected_project['id']
    selected_lot = st.session_state.get('selected_lot')
    
    if 'generated_structure' not in st.session_state or not st.session_state.generated_structure:
        st.warning("No se ha generado ninguna estructura. Por favor, vuelve al paso anterior.")
        if st.button("← Volver a Fase 2"): go_to_phase2(); st.rerun()
        return

    if selected_lot:
        if selected_lot == OPCION_ANALISIS_GENERAL:
            st.success("🎯 **Enfoque actual:** Se realizará un análisis general para todo el proyecto.")
        else:
            st.success(f"🎯 **Enfoque actual:** Lote / Bloque: **{selected_lot}**")
    
    st.button("← Volver a la gestión de archivos (Fase 2)", on_click=go_to_phase2)

    def handle_regeneration_with_feedback():
        feedback_text = st.session_state.get("feedback_area", "")
        if not feedback_text.strip():
            st.warning("Por favor, escribe tus indicaciones en el área de texto para la regeneración.")
            return

        with st.spinner("🧠 Incorporando tu feedback y regenerando la estructura..."):
            try:
                idioma_seleccionado = st.session_state.get('project_language', 'Español')
                contexto_lote = get_lot_context()
                prompt_con_idioma = PROMPT_REGENERACION.format(idioma=idioma_seleccionado, contexto_lote=contexto_lote)
                
                contenido_ia_regeneracion = [
                    prompt_con_idioma,
                    "--- INSTRUCCIONES DEL USUARIO ---\n" + feedback_text,
                    "--- ESTRUCTURA JSON ANTERIOR A CORREGIR ---\n" + json.dumps(st.session_state.generated_structure, indent=2, ensure_ascii=False)
                ]
                
                if st.session_state.get('uploaded_pliegos'):
                    st.write("Analizando documentos de referencia para la regeneración...")
                    
                    for file_info in st.session_state.uploaded_pliegos:
                        file_content_bytes = download_file_from_drive_cached(service, file_info['id'])
                        nombre_archivo = file_info['name']
                        
                        if nombre_archivo.lower().endswith('.xlsx'):
                            texto_csv = convertir_excel_a_texto_csv(file_content_bytes, nombre_archivo)
                            if texto_csv: contenido_ia_regeneracion.append(texto_csv)
                        else:
                            contenido_ia_regeneracion.append({
                                "mime_type": file_info['mimeType'], 
                                "data": file_content_bytes.getvalue()
                            })

                generation_config = genai.GenerationConfig(response_mime_type="application/json")
                response_regeneracion = model.generate_content(contenido_ia_regeneracion, generation_config=generation_config)
                
                if not response_regeneracion.candidates: st.error("La IA no generó una respuesta."); return

                json_limpio_str_regenerado = limpiar_respuesta_json(response_regeneracion.text)
                
                if json_limpio_str_regenerado:
                    st.session_state.generated_structure = json.loads(json_limpio_str_regenerado)
                    st.toast("¡Estructura regenerada con tu feedback!")
                    st.session_state.feedback_area = ""
                    st.rerun()
                else:
                    st.error("La IA no devolvió una estructura JSON válida tras la regeneración.")

            except Exception as e:
                st.error(f"Ocurrió un error crítico durante la regeneración: {e}")

    with st.container(border=True):
        st.subheader("Índice Propuesto y Guía de Redacción")
        estructura = st.session_state.generated_structure.get('estructura_memoria')
        matices = st.session_state.generated_structure.get('matices_desarrollo')
        mostrar_indice_desplegable(estructura, matices)
        
        st.markdown("---")
        st.subheader("📊 Plan Estratégico del Documento")
        config = st.session_state.generated_structure.get('configuracion_licitacion', {})
        plan = st.session_state.generated_structure.get('plan_extension', [])
        
        if config or plan:
            col1, col2 = st.columns(2)
            with col1: st.metric("Páginas Máximas", config.get('max_paginas', 'N/D'))
            with col2: st.metric("Reglas de Formato", config.get('reglas_formato', 'N/D'))

            st.markdown("---")
            if plan:
                try:
                    plan_data = []
                    for item in plan:
                        plan_data.append({
                            'Apartado': item.get('apartado', 'Sin Título'),
                            'Páginas Sugeridas': item.get('paginas_sugeridas_apartado', 'N/D'),
                            'Puntuación': item.get('puntuacion_sugerida', 'N/D')
                        })
                    df = pd.DataFrame(plan_data)
                    st.write("Distribución de Contenido y Puntuación:")
                    st.dataframe(df, use_container_width=True)
                except Exception as e: st.error(f"No se pudo mostrar el plan de extensión. Error: {e}")
            else: st.info("No se encontró un 'plan_extension' en la estructura generada.")
        else: st.warning("No se encontraron datos de 'configuracion_licitacion' o 'plan_extension' en la estructura generada por la IA.")

    st.markdown("---")
    st.subheader("Validación y Siguiente Paso")
    
    st.text_area(
        "Si necesitas cambios en el índice, el plan o las indicaciones, descríbelos aquí:",
        key="feedback_area",
        placeholder="Ejemplos:\n- 'El límite real son 40 páginas, reajusta la distribución.'\n- 'En el apartado 2, une los subapartados 2.1 y 2.2.'"
    )
    
    col1, col2 = st.columns(2)
    with col1:
        st.button("Regenerar con Feedback", on_click=handle_regeneration_with_feedback, use_container_width=True)
    with col2:
        st.button("🔁 Regenerar Todo desde Cero", on_click=lambda: handle_full_regeneration(model), use_container_width=True, help="Descarta este análisis y genera uno nuevo leyendo los archivos desde cero.")

    if st.button("Aceptar y Pasar a Fase 3 →", type="primary", use_container_width=True):
        with st.spinner("Guardando análisis final y preparando carpetas..."):
            try:
                index_folder_id, index_filename = get_lot_index_info(service, project_folder_id, selected_lot)
                json_bytes = json.dumps(st.session_state.generated_structure, indent=2, ensure_ascii=False).encode('utf-8')
                mock_file_obj = io.BytesIO(json_bytes)
                mock_file_obj.name = index_filename
                mock_file_obj.type = "application/json"
                
                saved_index_id = find_file_by_name(service, index_filename, index_folder_id)
                if saved_index_id:
                    delete_file_from_drive(service, saved_index_id)
                
                upload_file_to_drive(service, mock_file_obj, index_folder_id)
                st.toast(f"Análisis final guardado como '{index_filename}' en tu Drive.")
                
                st.toast("Creando estructura de carpetas para los guiones...")
                active_lot_folder_id = get_or_create_lot_folder_id(service, project_folder_id, lot_name=selected_lot)
                sync_guiones_folders_with_index(service, active_lot_folder_id, st.session_state.generated_structure)
                st.toast("¡Estructura de carpetas lista para la Fase 3!")

                go_to_phase3()
                st.rerun()
            except Exception as e:
                st.error(f"Ocurrió un error durante el guardado y la creación de carpetas: {e}")

# =============================================================================
#           FASE 3: CENTRO DE MANDO DE GUIONES
# =============================================================================

def ejecutar_generacion_con_gemini(model, credentials, project_folder_id, active_lot_folder_id, titulo, indicaciones_completas, contexto_adicional_lotes="", project_language='Español'):
    """
    (VERSIÓN MEJORADA)
    Genera el guion para un subapartado. Ahora analiza correctamente los archivos .docx
    de contexto antes de enviarlos a la IA, evitando el error de MIME type.
    """
    from googleapiclient.discovery import build
    service = build('drive', 'v3', credentials=credentials)

    nombre_limpio = clean_folder_name(titulo)
    nombre_archivo = nombre_limpio + ".docx"
    
    try:
        guiones_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=active_lot_folder_id)
        subapartado_guion_folder_id = find_or_create_folder(service, nombre_limpio, parent_id=guiones_folder_id)
        pliegos_folder_id = find_or_create_folder(service, "Pliegos", parent_id=project_folder_id)
        
        contexto_lote_actual = get_lot_context()
        prompt = PROMPT_GEMINI_PROPUESTA_ESTRATEGICA.format(idioma=project_language, contexto_lote=contexto_lote_actual)
        
        contenido_ia = [prompt, "--- INDICACIONES PARA ESTE APARTADO ---\n" + json.dumps(indicaciones_completas, indent=2, ensure_ascii=False)]

        if contexto_adicional_lotes:
            contenido_ia.append(contexto_adicional_lotes)

        # --- Procesamiento de Pliegos con lógica multimodal ---
        pliegos_en_drive = get_files_in_project(service, pliegos_folder_id)
        for file_info in pliegos_en_drive:
            file_bytes_io = download_file_from_drive_uncached(service, file_info['id'])
            if 'wordprocessingml' in file_info['mimeType']:
                analisis_multimodal = analizar_docx_multimodal_con_gemini(file_bytes_io, file_info['name'])
                if analisis_multimodal and "Error" not in analisis_multimodal:
                    contenido_ia.append(analisis_multimodal)
            elif file_info['name'].lower().endswith('.xlsx'):
                texto_csv = convertir_excel_a_texto_csv(file_bytes_io, file_info['name'])
                if texto_csv: contenido_ia.append(texto_csv)
            else:
                contenido_ia.append({"mime_type": file_info['mimeType'], "data": file_bytes_io.getvalue()})

        # --- Procesamiento de Documentos de Apoyo con lógica multimodal (¡LA CORRECCIÓN CLAVE!) ---
        docs_de_apoyo = get_files_in_project(service, subapartado_guion_folder_id)
        docs_de_apoyo_filtrados = [f for f in docs_de_apoyo if not f['name'] == nombre_archivo]
        if docs_de_apoyo_filtrados:
            contenido_ia.append("\n--- DOCUMENTACIÓN DE APOYO ADICIONAL ---\n")
            for uploaded_file_info in docs_de_apoyo_filtrados:
                file_bytes_io_apoyo = download_file_from_drive_uncached(service, uploaded_file_info['id'])
                
                if 'wordprocessingml' in uploaded_file_info['mimeType']:
                    # ¡AQUÍ ESTÁ LA MEJORA! Si es un .docx, lo analizamos primero.
                    analisis_multimodal = analizar_docx_multimodal_con_gemini(file_bytes_io_apoyo, uploaded_file_info['name'])
                    if analisis_multimodal and "Error" not in analisis_multimodal:
                        contenido_ia.append(analisis_multimodal)
                
                elif uploaded_file_info['name'].lower().endswith('.xlsx'):
                    texto_csv = convertir_excel_a_texto_csv(file_bytes_io_apoyo, uploaded_file_info['name'])
                    if texto_csv: contenido_ia.append(texto_csv)
                
                else:
                    # Para otros tipos de archivo soportados (como PDF), los enviamos directamente.
                    contenido_ia.append({"mime_type": uploaded_file_info['mimeType'], "data": file_bytes_io_apoyo.getvalue()})
        
        # El resto de la función sigue igual
        chat = model.start_chat()
        response = enviar_mensaje_con_reintentos(chat, contenido_ia)
        if not response or not response.candidates:
            # Imprime el feedback de la API si la respuesta fue bloqueada
            if response and hasattr(response, 'prompt_feedback'):
                st.error(f"La generación para '{titulo}' fue bloqueada. Razón: {response.prompt_feedback}")
            else:
                 st.error(f"No se obtuvo respuesta válida de la API para '{titulo}'.")
            return False
        
        documento = docx.Document()
        agregar_markdown_a_word(documento, response.text)
        doc_io = io.BytesIO()
        documento.save(doc_io); doc_io.seek(0)
        
        word_file_obj = io.BytesIO(doc_io.getvalue())
        word_file_obj.name = nombre_archivo
        word_file_obj.type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        existing_guion = find_file_by_name(service, nombre_archivo, subapartado_guion_folder_id)
        if existing_guion: delete_file_from_drive(service, existing_guion)
        
        upload_file_to_drive(service, word_file_obj, subapartado_guion_folder_id)
        return True

    except Exception as e:
        # Captura y muestra errores específicos de la API de Gemini si están disponibles
        if isinstance(e, genai.types.StopCandidateException):
            st.error(f"Error en la generación para '{titulo}': La respuesta fue detenida prematuramente. Esto puede deberse a los filtros de seguridad. {e}")
        else:
            st.error(f"Error inesperado en el hilo de generación para '{titulo}': {e}")
        # Imprime en la consola del servidor para depuración
        print(f"ERROR en el hilo de generación para '{titulo}': {e}")
        return False

def ejecutar_fase_4_en_background(model, credentials, project_folder_id, active_lot_folder_id, all_matices, generated_structure, project_language):
    """
    Esta función orquesta toda la lógica de la Fase 4:
    1. Genera todos los planes de prompts individuales en paralelo.
    2. Unifica los resultados en un solo archivo JSON.
    Devuelve True si todo fue exitoso, False si algo falló.
    """
    from googleapiclient.discovery import build
    service = build('drive', 'v3', credentials=credentials)
    
    st.info("Iniciando la generación de todos los planes de prompts en paralelo...")
    MAX_WORKERS = 4
    completed_count = 0
    all_successful = True
    
    # --- Parte 1: Generación en Paralelo de Prompts Individuales ---
    guiones_main_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=active_lot_folder_id)
    carpetas_de_guiones = list_project_folders(service, guiones_main_folder_id)
    
    items_to_generate = [
        matiz for matiz in all_matices 
        if clean_folder_name(matiz.get('subapartado')) in carpetas_de_guiones
    ]
    num_items = len(items_to_generate)
    
    if num_items == 0:
        st.warning("No se encontraron guiones generados para crear planes de prompts. Asegúrate de generar los borradores primero.")
        return False

    with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        future_to_matiz = {
            executor.submit(
                ejecutar_generacion_prompts_en_hilo, 
                model, credentials, project_folder_id, active_lot_folder_id,
                matiz, generated_structure, project_language
            ): matiz for matiz in items_to_generate
        }
        
        for future in concurrent.futures.as_completed(future_to_matiz):
            matiz_info = future_to_matiz[future]
            titulo = matiz_info.get('subapartado', 'Desconocido')
            try:
                success = future.result()
                if not success:
                    st.error(f"❌ Falló la generación del plan para: {titulo}")
                    all_successful = False
            except Exception as exc:
                st.error(f"❌ Error crítico generando plan para '{titulo}': {exc}")
                all_successful = False
            
            completed_count += 1
            st.info(f"Planes de prompts completados: {completed_count}/{num_items} - {titulo}")

    if not all_successful:
        st.error("Algunos planes de prompts no se pudieron generar. El proceso se detendrá.")
        return False

    st.success("Todos los planes individuales han sido generados. Procediendo a la unificación...")

    # --- Parte 2: Unificación de todos los JSONs ---
    try:
        plan_conjunto_final = {"plan_de_prompts": []}
        carpetas_de_guiones_actualizadas = list_project_folders(service, guiones_main_folder_id)

        for nombre_carpeta, folder_id in carpetas_de_guiones_actualizadas.items():
            plan_id = find_file_by_name(service, "prompts_individual.json", folder_id)
            if plan_id:
                json_bytes = download_file_from_drive_uncached(service, plan_id).getvalue()
                plan_individual_obj = json.loads(json_bytes.decode('utf-8'))
                prompts_de_este_plan = plan_individual_obj.get("plan_de_prompts", [])
                plan_conjunto_final["plan_de_prompts"].extend(prompts_de_este_plan)
        
        if not plan_conjunto_final["plan_de_prompts"]:
            st.warning("No se encontraron planes para unificar, aunque la generación pareció exitosa.")
            return False
            
        lot_docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=active_lot_folder_id)
        lot_name_clean = clean_folder_name(st.session_state.selected_lot)
        nombre_archivo_final = f"plan_de_prompts_{lot_name_clean}.json"
        
        json_bytes_finales = json.dumps(plan_conjunto_final, indent=2, ensure_ascii=False).encode('utf-8')
        mock_file_obj = io.BytesIO(json_bytes_finales); mock_file_obj.name = nombre_archivo_final; mock_file_obj.type = "application/json"
        
        old_conjunto_id = find_file_by_name(service, nombre_archivo_final, lot_docs_app_folder_id)
        if old_conjunto_id: delete_file_from_drive(service, old_conjunto_id)
        
        upload_file_to_drive(service, mock_file_obj, lot_docs_app_folder_id)
        st.success(f"¡Plan conjunto para '{st.session_state.selected_lot}' generado! Total: {len(plan_conjunto_final['plan_de_prompts'])} prompts.")
        return True
    except Exception as e:
        st.error(f"Ocurrió un error crítico durante la unificación de los planes: {e}")
        return False

def phase_3_page(model, go_to_phase2_results, go_to_phase4):
    st.markdown("<h3>FASE 3: Centro de Mando de Guiones</h3>", unsafe_allow_html=True)
    st.markdown("Gestiona tus guiones y documentos de apoyo. Cuando termines, avanza a la siguiente fase para preparar la redacción.")
    st.markdown("---")

    # --- 1. Inicialización y Verificación de Sesión ---
    if 'regenerating_item' not in st.session_state: st.session_state.regenerating_item = None
    if 'uploader_key' not in st.session_state: st.session_state.uploader_key = 0
    if 'classification_results' not in st.session_state: st.session_state.classification_results = []

    service = st.session_state.drive_service
    project_folder_id = st.session_state.selected_project['id']
    selected_lot = st.session_state.get('selected_lot')
    active_lot_folder_id = get_or_create_lot_folder_id(service, project_folder_id, lot_name=selected_lot)

    if not active_lot_folder_id:
        st.warning("No se puede continuar sin un lote seleccionado. Vuelve a la Fase 1."); return

    index_folder_id, index_filename = get_lot_index_info(service, project_folder_id, selected_lot)

    if 'generated_structure' not in st.session_state:
        st.info(f"Sincronizando índice ('{index_filename}') desde Google Drive...")
        try:
            saved_index_id = find_file_by_name(service, index_filename, index_folder_id)
            if saved_index_id:
                index_content_bytes = download_file_from_drive_cached(service, saved_index_id)
                st.session_state.generated_structure = json.loads(index_content_bytes.getvalue().decode('utf-8'))
                st.rerun()
            else:
                st.warning(f"No se ha encontrado un índice guardado ('{index_filename}').")
                if st.button("← Ir a Fase 2"): go_to_phase2_results(); st.rerun()
                return
        except Exception as e:
            st.error(f"Error al cargar el índice desde Drive: {e}"); return

    sync_guiones_folders_with_index(service, active_lot_folder_id, st.session_state.generated_structure)

    # --- 2. Preparación de datos para la UI ---
    estructura = st.session_state.generated_structure.get('estructura_memoria', [])
    matices_originales = st.session_state.generated_structure.get('matices_desarrollo', [])
    subapartados_a_mostrar = []
    
    if any(seccion.get('subapartados') for seccion in estructura):
        for seccion in estructura:
            for subapartado_titulo in seccion.get('subapartados', []):
                matiz = next((m for m in matices_originales if m.get('subapartado') == subapartado_titulo), None)
                if matiz: subapartados_a_mostrar.append(matiz)
    else:
        for seccion in estructura:
            apartado_titulo = seccion.get('apartado')
            if apartado_titulo: subapartados_a_mostrar.append({"apartado": apartado_titulo, "subapartado": apartado_titulo, "indicaciones": f"Generar guion para {apartado_titulo}"})

    # --- 3. Lógica de Clasificación y Subida Automática de Contexto ---
    st.subheader("Central de Documentos de Contexto")
    # ... (Esta sección no cambia)
    with st.container(border=True):
        st.info("Sube aquí TODOS los documentos de apoyo (PDFs, Word con imágenes, etc.). La IA los analizará y asignará al subapartado correcto automáticamente.")
        context_files = st.file_uploader(
            "Arrastra aquí tus archivos de contexto",
            type=['pdf', 'docx', 'xlsx'],
            accept_multiple_files=True,
            key=f"central_context_uploader_{st.session_state.uploader_key}"
        )
        if st.button("🤖 Clasificar y Asignar Documentos", disabled=not context_files, type="primary"):
            if context_files:
                st.session_state.classification_results = []
                lista_titulos_subapartados = [matiz.get('subapartado') for matiz in subapartados_a_mostrar]
                json_titulos = json.dumps(lista_titulos_subapartados, ensure_ascii=False)
                progress_bar = st.progress(0, text="Iniciando clasificación...")
                status_placeholder = st.empty()

                for i, file in enumerate(context_files):
                    file_name = file.name
                    progress_text = f"Procesando ({i+1}/{len(context_files)}): {file_name}"
                    progress_bar.progress((i + 1) / len(context_files), text=progress_text)
                    
                    with status_placeholder.container(border=True):
                        try:
                            file.seek(0)
                            file_bytes_io = io.BytesIO(file.getvalue())
                            mime_type = file.type
                            
                            contenido_para_gemini = [PROMPT_CLASIFICAR_DOCUMENTO]

                            if 'wordprocessingml' in mime_type:
                                analisis_multimodal = analizar_docx_multimodal_con_gemini(file_bytes_io, file_name)
                                if analisis_multimodal and "Error" not in analisis_multimodal:
                                    contenido_para_gemini.append("--- CONTENIDO DEL DOCUMENTO A CLASIFICAR (ANALIZADO) ---")
                                    contenido_para_gemini.append(analisis_multimodal)
                                else:
                                    st.error(f"No se pudo analizar el docx '{file_name}'. Se omitirá.")
                                    st.session_state.classification_results.append({"filename": file_name, "destination": "❌ Error de Análisis"})
                                    continue
                            else:
                                contenido_para_gemini.append("--- CONTENIDO DEL DOCUMENTO A CLASIFICAR (ORIGINAL) ---")
                                contenido_para_gemini.append({"mime_type": mime_type, "data": file.getvalue()})

                            contenido_para_gemini.append("--- ÍNDICE DE SUBAPARTADOS DISPONIBLES ---")
                            contenido_para_gemini.append(json_titulos)

                            st.write(f"Enviando contenido a la IA para clasificación...")
                            response = model.generate_content(
                                contenido_para_gemini,
                                generation_config={"response_mime_type": "application/json"}
                            )
                            
                            json_limpio = limpiar_respuesta_json(response.text)
                            resultado = json.loads(json_limpio)
                            subapartado_destino = resultado.get("subapartado_seleccionado")

                            if subapartado_destino and subapartado_destino != "inclasificable":
                                st.write(f"Destino: '{subapartado_destino}'. Subiendo a Google Drive...")
                                guiones_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=active_lot_folder_id)
                                nombre_limpio_carpeta = clean_folder_name(subapartado_destino)
                                destino_folder_id = find_or_create_folder(service, nombre_limpio_carpeta, parent_id=guiones_folder_id)
                                file.seek(0)
                                upload_file_to_drive(service, file, destino_folder_id)
                                st.success(f"✅ `{file_name}` asignado a **{subapartado_destino}**.")
                                st.session_state.classification_results.append({"filename": file_name, "destination": subapartado_destino})
                            else:
                                st.warning(f"⚠️ No se pudo clasificar `{file_name}`. Revisa si su contenido es relevante.")
                                st.session_state.classification_results.append({"filename": file_name, "destination": "⚠️ Inclasificable"})

                        except Exception as e:
                            st.error(f"Ocurrió un error procesando `{file_name}`: {e}")
                            st.session_state.classification_results.append({"filename": file_name, "destination": f"❌ Error Crítico"})
                
                progress_bar.empty(); status_placeholder.empty()
                st.toast("Proceso de clasificación finalizado.")
                st.session_state.uploader_key += 1 
                st.rerun()

    if st.session_state.classification_results:
        st.subheader("Resultados de la Última Clasificación")
        # ... (Esta sección no cambia)
        with st.container(border=True):
            df_results = pd.DataFrame(st.session_state.classification_results)
            df_results.rename(columns={'filename': 'Archivo', 'destination': 'Subapartado Asignado'}, inplace=True)
            st.dataframe(df_results, use_container_width=True, hide_index=True)
            if st.button("Limpiar resultados", key="clear_results"):
                st.session_state.classification_results = []; st.rerun()

    # --- 4. Funciones de Lógica Interna (Callbacks) ---
    # ... (Estas funciones de handle_*, ejecutar_*, etc., no cambian)
    def handle_confirm_regeneration(titulo, file_id_borrador, feedback):
        if not feedback.strip(): st.warning("Por favor, introduce tu feedback para la re-generación."); return
        with st.spinner(f"Re-generando '{titulo}' con tu feedback..."):
            try:
                borrador_bytes = download_file_from_drive_cached(service, file_id_borrador)
                doc = docx.Document(io.BytesIO(borrador_bytes.getvalue()))
                borrador_original_texto = "\n".join([p.text for p in doc.paragraphs])
                pliegos_folder_id = find_or_create_folder(service, "Pliegos", parent_id=project_folder_id)
                pliegos_en_drive = get_files_in_project(service, pliegos_folder_id)
                idioma = st.session_state.get('project_language', 'Español')
                contexto_lote_actual = get_lot_context()
                prompt = PROMPT_CONSULTOR_REVISION.format(idioma=idioma, contexto_lote=contexto_lote_actual)
                
                contenido_ia = [
                    prompt, 
                    "--- BORRADOR ORIGINAL ---\n" + borrador_original_texto, 
                    "--- FEEDBACK DEL CLIENTE (Tus correcciones y comentarios) ---\n" + feedback
                ]
                
                for file_info in pliegos_en_drive:
                    file_content_bytes = download_file_from_drive_cached(service, file_info['id'])
                    if 'wordprocessingml' in file_info['mimeType']:
                        analisis = analizar_docx_multimodal_con_gemini(file_content_bytes, file_info['name'])
                        if analisis: contenido_ia.append(analisis)
                    else:
                        contenido_ia.append({"mime_type": file_info['mimeType'], "data": file_content_bytes.getvalue()})

                response = model.generate_content(contenido_ia)
                if not response.candidates: st.error("La IA no generó una respuesta para la re-generación."); return
                
                documento_nuevo = docx.Document()
                agregar_markdown_a_word(documento_nuevo, response.text)
                doc_io = io.BytesIO(); documento_nuevo.save(doc_io)
                word_file_obj = io.BytesIO(doc_io.getvalue())
                nombre_limpio = clean_folder_name(titulo)
                nombre_archivo = nombre_limpio + ".docx"
                word_file_obj.name = nombre_archivo
                word_file_obj.type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                
                guiones_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=active_lot_folder_id)
                subapartado_guion_folder_id = find_or_create_folder(service, nombre_limpio, parent_id=guiones_folder_id)
                
                delete_file_from_drive(service, file_id_borrador)
                upload_file_to_drive(service, word_file_obj, subapartado_guion_folder_id)
                st.toast(f"¡Guion para '{titulo}' re-generado con éxito!")
                st.session_state.regenerating_item = None; st.rerun()
            except Exception as e:
                st.error(f"Error crítico durante la re-generación: {e}"); st.session_state.regenerating_item = None
    
    def ejecutar_regeneracion(titulo): st.session_state.regenerating_item = titulo
    def cancelar_regeneracion(): st.session_state.regenerating_item = None
    def ejecutar_borrado_de_guion_completo(titulo, folder_id_to_delete):
        with st.spinner(f"Eliminando guion y contexto para '{titulo}'..."):
            try:
                delete_file_from_drive(service, folder_id_to_delete); st.toast(f"Carpeta del guion '{titulo}' eliminada."); st.rerun()
            except Exception as e: st.error(f"Ocurrió un error inesperado al borrar la carpeta: {e}")
    def handle_context_file_delete(file_id, file_name):
        with st.spinner(f"Eliminando archivo '{file_name}'..."):
            try:
                delete_file_from_drive(service, file_id); st.toast(f"Archivo '{file_name}' eliminado."); st.rerun()
            except Exception as e: st.error(f"No se pudo eliminar el archivo: {e}")
    def handle_direct_context_upload(files, destination_folder_id):
        if files:
            with st.spinner(f"Subiendo {len(files)} archivo(s) de contexto..."):
                for file_obj in files: upload_file_to_drive(service, file_obj, destination_folder_id)
                st.toast("Archivos de contexto añadidos."); st.rerun()
    # --- 5. Renderizado de la Interfaz de Usuario ---
    st.markdown("---")
    st.subheader("Gestión de Guiones de Subapartados")
    # ... (Toda la lógica de la UI para mostrar los guiones, los botones de generar, borrar, etc., no cambia)
    with st.spinner("Sincronizando guiones y archivos de contexto con Google Drive..."):
        guiones_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=active_lot_folder_id)
        carpetas_existentes = {f['name']: f['id'] for f in get_files_in_project(service, guiones_folder_id) if f['mimeType'] == 'application/vnd.google-apps.folder'}
        guiones_generados_data = {}
        for nombre_carpeta, folder_id in carpetas_existentes.items():
            files_in_subfolder = get_files_in_project(service, folder_id)
            script_file = next((f for f in files_in_subfolder if f['name'] == nombre_carpeta + ".docx"), None)
            context_files_list = [f for f in files_in_subfolder if f != script_file]
            guiones_generados_data[nombre_carpeta] = {"script": script_file, "context": context_files_list, "folder_id": folder_id}

    pending_keys = [matiz.get('subapartado') for matiz in subapartados_a_mostrar if clean_folder_name(matiz.get('subapartado')) not in guiones_generados_data or not guiones_generados_data[clean_folder_name(matiz.get('subapartado'))]['script']]
    
    def toggle_all_checkboxes():
        new_state = st.session_state.get('select_all_checkbox', False)
        for key in pending_keys: st.session_state[f"cb_{key}"] = new_state

    with st.container(border=True):
        st.subheader("Generación de Borradores en Lote (Paralelo)")
        col_sel_1, col_sel_2 = st.columns([1, 2])
        with col_sel_1:
            st.checkbox("Seleccionar Todos / Ninguno", key="select_all_checkbox", on_change=toggle_all_checkboxes, disabled=not pending_keys)
        with col_sel_2:
            selected_keys = [key for key in pending_keys if st.session_state.get(f"cb_{key}")]
            num_selected = len(selected_keys)
            
            if st.button(f"🚀 Generar {num_selected} borradores en paralelo", type="primary", use_container_width=True, disabled=(num_selected == 0)):
                items_to_generate = [matiz for matiz in subapartados_a_mostrar if matiz.get('subapartado') in selected_keys]
                MAX_WORKERS = 4
                progress_bar = st.progress(0, text="Configurando generación en paralelo...")
                st.info(f"Se generarán {num_selected} guiones usando hasta {MAX_WORKERS} hilos. Esto puede tardar varios minutos.")
                completed_count = 0; all_successful = True

                credentials = get_credentials()
                project_language = st.session_state.get('project_language', 'Español')

                if not credentials:
                    st.error("Error de autenticación. No se puede iniciar la generación.")
                else:
                    with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
                        future_to_matiz = {
                            executor.submit(
                                ejecutar_generacion_con_gemini, 
                                model, credentials, project_folder_id, active_lot_folder_id,
                                matiz.get('subapartado'), matiz, "", project_language
                            ): matiz for matiz in items_to_generate
                        }
                        for future in concurrent.futures.as_completed(future_to_matiz):
                            matiz_info = future_to_matiz[future]
                            titulo = matiz_info.get('subapartado', 'Desconocido')
                            try:
                                success = future.result()
                                if not success:
                                    st.error(f"❌ Falló la generación para: {titulo}")
                                    all_successful = False
                            except Exception as exc:
                                st.error(f"❌ Error crítico generando '{titulo}': {exc}")
                                all_successful = False
                            completed_count += 1
                            progress_text = f"Completados {completed_count}/{num_selected}: {titulo}"
                            progress_bar.progress(completed_count / num_selected, text=progress_text)

                if all_successful:
                    progress_bar.progress(1.0, text="¡Generación en lote completada!"); st.success(f"{num_selected} borradores generados."); st.balloons()
                else:
                    st.warning("Algunos guiones no se pudieron generar. Revisa la consola para más detalles.")
                time.sleep(4); st.rerun()

    st.markdown("---")
    for i, matiz in enumerate(subapartados_a_mostrar):
        subapartado_titulo = matiz.get('subapartado'); 
        if not subapartado_titulo: continue
        
        nombre_limpio = clean_folder_name(subapartado_titulo)
        drive_data = guiones_generados_data.get(nombre_limpio)
        guion_generado = drive_data and drive_data.get("script")
        estado = "📄 Generado" if guion_generado else "⚪ No Generado"
        file_info = drive_data.get("script") if drive_data else None
        subapartado_folder_id = drive_data.get("folder_id") if drive_data else None
        context_files_list = drive_data.get("context", []) if drive_data else []
        
        with st.container(border=True):
            col1, col2 = st.columns([2, 1])
            with col1:
                if estado == "⚪ No Generado": st.checkbox(f"**{subapartado_titulo}**", key=f"cb_{subapartado_titulo}")
                else: st.write(f"**{subapartado_titulo}**")
                st.caption(f"Estado: {estado}")

                with st.expander(f"Gestionar ({len(context_files_list)}) archivos de contexto"):
                    if not subapartado_folder_id:
                        if st.button("Crear carpeta para añadir contexto", key=f"create_folder_{i}"):
                            guiones_main_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=active_lot_folder_id)
                            find_or_create_folder(service, nombre_limpio, parent_id=guiones_main_folder_id)
                            st.rerun()
                    else:
                        if not context_files_list: st.info("No hay archivos de contexto asignados a este subapartado.")
                        for ctx_file in context_files_list:
                            c1, c2 = st.columns([3, 1])
                            c1.write(f"📄 **{ctx_file['name']}**")
                            c2.button("Eliminar", key=f"del_ctx_{ctx_file['id']}", on_click=handle_context_file_delete, args=(ctx_file['id'], ctx_file['name']))
                        st.markdown("---")
                        new_context_files = st.file_uploader("Añadir nuevos archivos de contexto aquí:", accept_multiple_files=True, key=f"uploader_ctx_{i}")
                        st.button("Guardar en Drive", key=f"upload_ctx_{i}", on_click=handle_direct_context_upload, args=(new_context_files, subapartado_folder_id), type="primary", disabled=not new_context_files)

                if st.session_state.regenerating_item == subapartado_titulo and file_info:
                    st.text_area("Introduce tus correcciones o feedback para mejorar este guion:", key=f"feedback_{i}", height=150)
                    c1, c2 = st.columns(2)
                    c1.button("✅ Confirmar Re-generación", key=f"confirm_regen_{i}", on_click=handle_confirm_regeneration, args=(subapartado_titulo, file_info['id'], st.session_state[f"feedback_{i}"]), type="primary")
                    c2.button("❌ Cancelar", key=f"cancel_regen_{i}", on_click=cancelar_regeneracion)

            with col2:
                if estado == "📄 Generado" and file_info:
                    st.link_button("Revisar Guion", f"https://docs.google.com/document/d/{file_info['id']}/edit", use_container_width=True)
                    st.button("Re-Generar con Feedback", key=f"regen_{i}", on_click=ejecutar_regeneracion, args=(subapartado_titulo,), type="secondary", use_container_width=True)
                    st.button("🗑️ Borrar Carpeta", key=f"del_{i}", on_click=ejecutar_borrado_de_guion_completo, args=(subapartado_titulo, subapartado_folder_id), use_container_width=True)
                else:
                    if st.button("Generar Borrador", key=f"gen_{i}", use_container_width=True):
                        with st.spinner(f"Generando borrador para '{subapartado_titulo}'..."):
                            credentials = get_credentials()
                            project_language = st.session_state.get('project_language', 'Español')
                            if not credentials:
                                st.error("Error de autenticación. Por favor, reinicia la sesión.")
                            else:
                                success = ejecutar_generacion_con_gemini(
                                    model=model, credentials=credentials,
                                    project_folder_id=project_folder_id, active_lot_folder_id=active_lot_folder_id,
                                    titulo=subapartado_titulo, indicaciones_completas=matiz, project_language=project_language
                                )
                                if success: st.rerun()

    # --- 6. NAVEGACIÓN DE PÁGINA (LÓGICA SIMPLIFICADA) ---
    st.markdown("---")
    
    with st.container(border=True):
        st.subheader("Finalizar Fase y Avanzar")
        st.info("Una vez que hayas generado y revisado todos los guiones necesarios, avanza a la siguiente fase para preparar los prompts detallados para la redacción final.")
        
        # Este botón ahora solo se encarga de la navegación a la Fase 4.
        # La lógica de 'on_click' se encarga de cambiar la página en st.session_state.
        st.button("Avanzar a Fase 4 (Preparación de Prompts) →", 
                        on_click=go_to_phase4, 
                        type="primary", 
                        use_container_width=True)

    st.markdown("---")
    st.button("← Volver a Revisión de Índice (F2)", on_click=go_to_phase2_results, use_container_width=True)

# =============================================================================
#           FASE 4: CENTRO DE MANDO DE PROMPTS
# =============================================================================

# ----------------- ¡NUEVA FUNCIÓN TRABAJADORA (WORKER)! -----------------
def ejecutar_generacion_prompts_en_hilo(model, credentials, project_folder_id, active_lot_folder_id, matiz_info, generated_structure_dict, project_language):
    """
    Función segura para hilos que genera un plan de prompts para un subapartado.
    """
    from googleapiclient.discovery import build
    service = build('drive', 'v3', credentials=credentials)

    apartado_titulo = matiz_info.get("apartado", "N/A")
    subapartado_titulo = matiz_info.get("subapartado", "N/A")
    
    try:
        config_licitacion = generated_structure_dict.get('configuracion_licitacion', {})
        plan_extension = generated_structure_dict.get('plan_extension', [])
        max_paginas_doc = config_licitacion.get('max_paginas', 'N/D')
        reglas_formato_doc = config_licitacion.get('reglas_formato', 'N/D')

        min_chars_sub, max_chars_sub, paginas_sugeridas_sub = 3500, 3800, "1"
        for item_apartado in plan_extension:
            if item_apartado.get('apartado') == apartado_titulo:
                for item_subapartado in item_apartado.get('desglose_subapartados', []):
                    if item_subapartado.get('subapartado') == subapartado_titulo:
                        min_chars_sub = item_subapartado.get('min_caracteres_sugeridos', min_chars_sub)
                        max_chars_sub = item_subapartado.get('max_caracteres_sugeridos', max_chars_sub)
                        paginas_sugeridas_sub = str(item_subapartado.get('paginas_sugeridas', paginas_sugeridas_sub))
                        break
                break
        
        guiones_main_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=active_lot_folder_id)
        nombre_limpio = clean_folder_name(subapartado_titulo)
        subapartado_folder_id = find_or_create_folder(service, nombre_limpio, parent_id=guiones_main_folder_id)
        
        contexto_adicional_str = ""
        files_in_subfolder = get_files_in_project(service, subapartado_folder_id)
        
        for file_info in files_in_subfolder:
            if file_info['name'].lower().endswith('.docx'):
                # Usar la versión sin caché para seguridad en hilos
                file_bytes = download_file_from_drive_uncached(service, file_info['id'])
                doc = docx.Document(io.BytesIO(file_bytes.getvalue()))
                texto_doc = "\n".join([p.text for p in doc.paragraphs])
                contexto_adicional_str += f"\n--- CONTENIDO DEL GUION ({file_info['name']}) ---\n{texto_doc}\n"
        
        prompt_final = PROMPT_DESARROLLO.format(
            idioma=project_language, max_paginas=max_paginas_doc,
            reglas_formato=reglas_formato_doc, apartado_referencia=apartado_titulo,
            subapartado_referencia=subapartado_titulo, paginas_sugeridas_subapartado=paginas_sugeridas_sub,
            min_chars_total=min_chars_sub, max_chars_total=max_chars_sub
        )
        
        contenido_ia = [prompt_final]
        if contexto_adicional_str:
            contenido_ia.append(contexto_adicional_str)
        else:
            print(f"ADVERTENCIA (hilo): No se encontró guion para '{subapartado_titulo}'.")

        response = model.generate_content(contenido_ia, generation_config={"response_mime_type": "application/json"})
        json_limpio_str = limpiar_respuesta_json(response.text)
        
        if json_limpio_str:
            plan_parcial_obj = json.loads(json_limpio_str)
            json_bytes = json.dumps(plan_parcial_obj, indent=2, ensure_ascii=False).encode('utf-8')
            mock_file_obj = io.BytesIO(json_bytes); mock_file_obj.name = "prompts_individual.json"; mock_file_obj.type = "application/json"
            
            old_plan_id = find_file_by_name(service, "prompts_individual.json", subapartado_folder_id)
            if old_plan_id: delete_file_from_drive(service, old_plan_id)
            upload_file_to_drive(service, mock_file_obj, subapartado_folder_id)
            return True
        return False
    except Exception as e:
        print(f"ERROR en el hilo de generación de prompts para '{subapartado_titulo}': {e}")
        return False

def phase_4_page(model, go_to_phase3, go_to_phase5):
    st.markdown("<h3>FASE 4: Centro de Mando de Prompts</h3>", unsafe_allow_html=True)
    st.markdown("Genera planes de prompts de forma individual o selecciónalos para procesarlos en lote.")
    st.markdown("---")
    
    service = st.session_state.drive_service
    project_folder_id = st.session_state.selected_project['id']
    selected_lot = st.session_state.get('selected_lot')
    active_lot_folder_id = get_or_create_lot_folder_id(service, project_folder_id, lot_name=selected_lot)
    
    if not active_lot_folder_id:
        st.warning("No se puede continuar sin un lote seleccionado. Vuelve a la Fase 1."); return

    index_folder_id, index_filename = get_lot_index_info(service, project_folder_id, selected_lot)

    if 'generated_structure' not in st.session_state:
        st.info(f"Sincronizando índice ('{index_filename}') desde Google Drive...")
        try:
            saved_index_id = find_file_by_name(service, index_filename, index_folder_id)
            if saved_index_id:
                index_content_bytes = download_file_from_drive_cached(service, saved_index_id)
                st.session_state.generated_structure = json.loads(index_content_bytes.getvalue().decode('utf-8'))
                st.rerun()
            else:
                st.warning(f"No se ha encontrado un índice guardado ('{index_filename}') para este lote. Vuelve a Fase 2 para generarlo.")
                if st.button("← Ir a Fase 2"): go_to_phase2_results(); st.rerun()
                return
        except Exception as e:
            st.error(f"Error al cargar el índice desde Drive: {e}"); return

    estructura = st.session_state.generated_structure.get('estructura_memoria', [])
    matices_originales = st.session_state.generated_structure.get('matices_desarrollo', [])
    matices_dict = {item.get('subapartado', ''): item for item in matices_originales if isinstance(item, dict) and 'subapartado' in item}
    if not estructura: st.error("La estructura JSON no contiene la clave 'estructura_memoria'."); return

    subapartados_a_mostrar = []
    hay_subapartados = any(seccion.get('subapartados') for seccion in estructura)
    if hay_subapartados:
        for seccion in estructura:
            apartado_principal = seccion.get('apartado', 'Sin Título')
            for subapartado_titulo in seccion.get('subapartados', []):
                matiz_existente = matices_dict.get(subapartado_titulo)
                if matiz_existente: subapartados_a_mostrar.append(matiz_existente)
                else: subapartados_a_mostrar.append({"apartado": apartado_principal, "subapartado": subapartado_titulo, "indicaciones": "No se encontraron indicaciones detalladas."})
    else:
        st.info("El índice no contiene subapartados. Se mostrarán los apartados principales.")
        for seccion in estructura:
            apartado_titulo = seccion.get('apartado')
            if apartado_titulo: subapartados_a_mostrar.append({"apartado": apartado_titulo, "subapartado": apartado_titulo, "indicaciones": f"Generar prompts para: {apartado_titulo}"})
    if not subapartados_a_mostrar: st.warning("El índice está vacío o tiene un formato incorrecto."); return

    def handle_individual_generation(matiz_info, show_toast=True):
        credentials = get_credentials()
        project_language = st.session_state.get('project_language', 'Español')
        if not credentials:
            st.error("Error de autenticación. No se puede proceder.")
            return False
            
        success = ejecutar_generacion_prompts_en_hilo(
            model, credentials, project_folder_id, active_lot_folder_id, 
            matiz_info, st.session_state.generated_structure, project_language
        )
        if success:
            if show_toast: st.toast(f"Plan para '{matiz_info.get('subapartado')}' generado.")
            st.rerun()
        else:
            st.error(f"Falló la generación del plan para '{matiz_info.get('subapartado')}'.")

    def handle_individual_deletion(titulo, plan_id_to_delete):
        with st.spinner(f"Eliminando el plan para '{titulo}'..."):
            if delete_file_from_drive(service, plan_id_to_delete):
                st.toast(f"Plan para '{titulo}' eliminado con éxito."); st.rerun()

    def handle_conjunto_generation():
        with st.spinner("Unificando todos los planes de prompts..."):
            try:
                guiones_main_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=active_lot_folder_id)
                carpetas_de_guiones = list_project_folders(service, guiones_main_folder_id)
                
                plan_conjunto_final = {"plan_de_prompts": []}
                for nombre_carpeta, folder_id in carpetas_de_guiones.items():
                    plan_id = find_file_by_name(service, "prompts_individual.json", folder_id)
                    if plan_id:
                        json_bytes = download_file_from_drive_cached(service, plan_id).getvalue()
                        plan_individual_obj = json.loads(json_bytes.decode('utf-8'))
                        prompts_de_este_plan = plan_individual_obj.get("plan_de_prompts", [])
                        plan_conjunto_final["plan_de_prompts"].extend(prompts_de_este_plan)
                
                if not plan_conjunto_final["plan_de_prompts"]:
                    st.warning("No se encontraron planes individuales para unificar. Genera al menos uno."); return
                
                lot_docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=active_lot_folder_id)
                lot_name_clean = clean_folder_name(st.session_state.selected_lot)
                nombre_archivo_final = f"plan_de_prompts_{lot_name_clean}.json"
                
                json_bytes_finales = json.dumps(plan_conjunto_final, indent=2, ensure_ascii=False).encode('utf-8')
                mock_file_obj = io.BytesIO(json_bytes_finales); mock_file_obj.name = nombre_archivo_final; mock_file_obj.type = "application/json"
                
                old_conjunto_id = find_file_by_name(service, nombre_archivo_final, lot_docs_app_folder_id)
                if old_conjunto_id: delete_file_from_drive(service, old_conjunto_id)
                
                upload_file_to_drive(service, mock_file_obj, lot_docs_app_folder_id)
                st.success(f"¡Plan conjunto para '{st.session_state.selected_lot}' generado! Se unificaron {len(plan_conjunto_final['plan_de_prompts'])} prompts.")
                st.balloons()
            except Exception as e:
                st.error(f"Ocurrió un error durante la unificación: {e}")

    with st.spinner("Verificando estado de los planes de prompts..."):
        guiones_main_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=active_lot_folder_id)
        carpetas_de_guiones = list_project_folders(service, guiones_main_folder_id)
        planes_individuales_existentes = {}
        for nombre_carpeta, folder_id in carpetas_de_guiones.items():
            plan_id = find_file_by_name(service, "prompts_individual.json", folder_id)
            if plan_id: planes_individuales_existentes[nombre_carpeta] = plan_id

    st.subheader("Generación de Planes de Prompts en Lote")
    pending_keys = [
        matiz.get('subapartado') for matiz in subapartados_a_mostrar
        if clean_folder_name(matiz.get('subapartado')) in carpetas_de_guiones
        and clean_folder_name(matiz.get('subapartado')) not in planes_individuales_existentes
    ]

    def toggle_all_prompt_checkboxes():
        new_state = st.session_state.get('select_all_prompts_checkbox', False)
        for key in pending_keys: st.session_state[f"pcb_{key}"] = new_state

    with st.container(border=True):
        col_sel_1, col_sel_2 = st.columns([1, 2])
        with col_sel_1:
            st.checkbox("Seleccionar Todos / Ninguno", key="select_all_prompts_checkbox", on_change=toggle_all_prompt_checkboxes, disabled=not pending_keys)
        with col_sel_2:
            selected_keys = [key for key in pending_keys if st.session_state.get(f"pcb_{key}")]
            num_selected = len(selected_keys)
            
            # ----------------- ¡BLOQUE MODIFICADO PARA USAR WORKERS! -----------------
            if st.button(f"🚀 Generar {num_selected} planes en paralelo", type="primary", use_container_width=True, disabled=(num_selected == 0)):
                items_to_generate = [matiz for matiz in subapartados_a_mostrar if matiz.get('subapartado') in selected_keys]
                MAX_WORKERS = 4
                progress_bar = st.progress(0, text="Configurando generación en paralelo...")
                st.info(f"Se generarán {num_selected} planes de prompts usando hasta {MAX_WORKERS} hilos.")
                completed_count = 0
                all_successful = True
                
                credentials = get_credentials()
                project_language = st.session_state.get('project_language', 'Español')
                generated_structure_dict = st.session_state.generated_structure # Copiar el dict para pasarlo
                
                if not credentials:
                    st.error("Error de autenticación. No se puede iniciar la generación en paralelo.")
                else:
                    with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
                        future_to_matiz = {
                            executor.submit(
                                ejecutar_generacion_prompts_en_hilo, 
                                model, credentials, project_folder_id, active_lot_folder_id,
                                matiz, generated_structure_dict, project_language
                            ): matiz for matiz in items_to_generate
                        }
                        
                        for future in concurrent.futures.as_completed(future_to_matiz):
                            matiz_info = future_to_matiz[future]
                            titulo = matiz_info.get('subapartado', 'Desconocido')
                            try:
                                success = future.result()
                                if not success:
                                    st.error(f"❌ Falló la generación del plan para: {titulo}")
                                    all_successful = False
                            except Exception as exc:
                                st.error(f"❌ Error crítico generando plan para '{titulo}': {exc}")
                                all_successful = False
                            
                            completed_count += 1
                            progress_text = f"Completados {completed_count}/{num_selected}: {titulo}"
                            progress_bar.progress(completed_count / num_selected, text=progress_text)
                
                if all_successful:
                    progress_bar.progress(1.0, text="¡Generación en lote completada!")
                    st.success(f"{num_selected} planes de prompts generados.")
                    st.balloons()
                else:
                    st.warning("Algunos planes no se pudieron generar. Revisa la consola para más detalles.")
                
                time.sleep(4)
                st.rerun()

    st.markdown("---")
    st.subheader("Gestión de Planes de Prompts")

    for i, matiz in enumerate(subapartados_a_mostrar):
        subapartado_titulo = matiz.get("subapartado")
        if not subapartado_titulo: continue
        
        nombre_limpio = clean_folder_name(subapartado_titulo)
        guion_generado = nombre_limpio in carpetas_de_guiones
        plan_individual_id = planes_individuales_existentes.get(nombre_limpio)
        
        with st.container(border=True):
            col1, col2 = st.columns([2, 1])
            with col1:
                if not plan_individual_id and guion_generado:
                    st.checkbox(f"**{subapartado_titulo}**", key=f"pcb_{subapartado_titulo}")
                else:
                    st.write(f"**{subapartado_titulo}**")
                
                if not guion_generado:
                    st.warning("⚠️ Guion no generado en Fase 3. No se puede crear un plan.")
                elif plan_individual_id:
                    st.success("✔️ Plan generado")
                    with st.expander("Ver / Descargar Plan Individual"):
                        json_bytes = download_file_from_drive_cached(service, plan_individual_id).getvalue()
                        st.json(json_bytes.decode('utf-8'))
                        st.download_button("Descargar JSON", data=json_bytes, file_name=f"prompts_{nombre_limpio}.json", mime="application/json", key=f"dl_{i}")
                else:
                    st.info("⚪ Pendiente de generar plan de prompts")

            with col2:
                if not plan_individual_id:
                    st.button("Generar Plan de Prompts", key=f"gen_ind_{i}", on_click=handle_individual_generation, args=(matiz, True), use_container_width=True, type="primary", disabled=not guion_generado)
                else:
                    st.button("Re-generar Plan", key=f"gen_regen_{i}", on_click=handle_individual_generation, args=(matiz, True), use_container_width=True, type="secondary")
                    st.button("🗑️ Borrar Plan", key=f"del_plan_{i}", on_click=handle_individual_deletion, args=(subapartado_titulo, plan_individual_id), use_container_width=True)

    st.markdown("---")
    st.button("🚀 Unificar y Guardar Plan de Prompts para este Lote", on_click=handle_conjunto_generation, use_container_width=True, type="primary")
    col_nav3_1, col_nav3_2 = st.columns(2)
    with col_nav3_1:
        st.button("← Volver al Centro de Mando (F3)", on_click=go_to_phase3, use_container_width=True)
    with col_nav3_2:
        st.button("Ir a Redacción Final (F5) →", on_click=go_to_phase5, use_container_width=True)
        
# =============================================================================
#           PÁGINA FASE 5: REDACCIÓN DEL CUERPO DEL DOCUMENTO
# =============================================================================
def phase_5_page(model, go_to_phase4, go_to_phase6):
    st.markdown("<h3>FASE 5: Redacción del Cuerpo del Documento</h3>", unsafe_allow_html=True)
    st.markdown("Ejecuta el plan de prompts para generar el contenido completo de la memoria técnica.")
    st.markdown("---")
    
    service = st.session_state.drive_service
    project_folder_id = st.session_state.selected_project['id']

    selected_lot = st.session_state.get('selected_lot')
    if not selected_lot:
        st.warning("No se ha seleccionado un lote en la sesión. Por favor, vuelve a la Fase 1 para continuar.")
        return

    active_lot_folder_id = get_or_create_lot_folder_id(service, project_folder_id, selected_lot)
    docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=active_lot_folder_id)

    lot_name_clean = clean_folder_name(selected_lot)
    plan_filename = f"plan_de_prompts_{lot_name_clean}.json"
    plan_conjunto_id = find_file_by_name(service, plan_filename, docs_app_folder_id)

    if not plan_conjunto_id:
        st.warning(f"No se ha encontrado un plan de prompts ('{plan_filename}') para este lote. Vuelve a la Fase 4 para generarlo y unificarlo.")
        if st.button("← Ir a Fase 4"): 
            go_to_phase4()
            st.rerun()
        return

    try:
        json_bytes = download_file_from_drive_cached(service, plan_conjunto_id).getvalue()
        plan_de_accion = json.loads(json_bytes.decode('utf-8'))
        lista_de_prompts = plan_de_accion.get("plan_de_prompts", [])
        if lista_de_prompts:
            lista_de_prompts.sort(key=lambda x: natural_sort_key(x.get('subapartado_referencia', '')))
        st.success(f"✔️ Plan de acción para '{selected_lot}' cargado. Se ejecutarán {len(lista_de_prompts)} prompts.")
    except Exception as e:
        st.error(f"Error al cargar o procesar el plan de acción: {e}")
        return

    button_text = "🔁 Volver a Generar Cuerpo del Documento" if st.session_state.get("generated_doc_buffer") else "🚀 Iniciar Redacción y Generar Cuerpo"
    
    if st.button(button_text, type="primary", use_container_width=True):
        if not lista_de_prompts:
            st.warning("El plan de acción está vacío. No hay nada que ejecutar.")
            return
            
        generation_successful = False
        documento = docx.Document()
        
        try:
            with st.spinner("Iniciando redacción... Esto puede tardar varios minutos."):
                fragment_counts = {}
                for tarea in lista_de_prompts:
                    sub_ref = tarea.get("subapartado_referencia")
                    if sub_ref: fragment_counts[sub_ref] = fragment_counts.get(sub_ref, 0) + 1
                
                character_budgets = {}
                plan_extension = st.session_state.generated_structure.get('plan_extension', [])
                for item_apartado in plan_extension:
                    for item_subapartado in item_apartado.get('desglose_subapartados', []):
                        sub_ref = item_subapartado.get('subapartado')
                        if sub_ref:
                            character_budgets[sub_ref] = (
                                item_subapartado.get('min_caracteres_sugeridos', 3500),
                                item_subapartado.get('max_caracteres_sugeridos', 3800)
                            )

                chat_redaccion = model.start_chat()
                progress_bar = st.progress(0, text="Configurando sesión de chat...")
                ultimo_apartado_escrito = None
                ultimo_subapartado_escrito = None

                for i, tarea in enumerate(lista_de_prompts):
                    subapartado_actual = tarea.get("subapartado_referencia")
                    apartado_actual = tarea.get("apartado_referencia")
                    progress_text = f"Procesando Tarea {i+1}/{len(lista_de_prompts)}: {subapartado_actual or 'N/A'}"
                    progress_bar.progress((i + 1) / len(lista_de_prompts), text=progress_text)
                    
                    if apartado_actual and apartado_actual != ultimo_apartado_escrito:
                        if ultimo_apartado_escrito is not None: documento.add_page_break()
                        documento.add_heading(apartado_actual, level=1)
                        ultimo_apartado_escrito = apartado_actual
                        ultimo_subapartado_escrito = None
                    
                    if subapartado_actual and subapartado_actual != ultimo_subapartado_escrito:
                        documento.add_heading(subapartado_actual, level=2)
                        ultimo_subapartado_escrito = subapartado_actual

                    prompt_actual = tarea.get("prompt_para_asistente")
                    respuesta_ia_bruta = ""
                    if prompt_actual:
                        prompt_a_enviar = prompt_actual
                        if subapartado_actual and '{min_chars_fragmento}' in prompt_actual:
                            num_fragments = fragment_counts.get(subapartado_actual, 1)
                            min_total, max_total = character_budgets.get(subapartado_actual, (3500, 3800))
                            min_per_fragment = min_total / num_fragments
                            max_per_fragment = max_total / num_fragments
                            prompt_a_enviar = prompt_actual.format(min_chars_fragmento=int(min_per_fragment), max_chars_fragmento=int(max_per_fragment))

                        response = enviar_mensaje_con_reintentos(chat_redaccion, prompt_a_enviar)
                        if not response:
                            st.error("La generación se ha detenido debido a un error persistente en la API."); generation_successful = False; break
                        respuesta_ia_bruta = response.text

                    es_html = ("HTML" in tarea.get("prompt_id", "").upper() or "VISUAL" in tarea.get("prompt_id", "").upper() or respuesta_ia_bruta.strip().startswith(('<!DOCTYPE html>', '<div', '<table')))
                    
                    if es_html:
                        html_puro = limpiar_respuesta_final(respuesta_ia_bruta)
                        image_file = html_a_imagen(wrap_html_fragment(html_puro), f"temp_img_{i}.png")
                        if image_file and os.path.exists(image_file):
                            documento.add_picture(image_file, width=docx.shared.Inches(6.5))
                            os.remove(image_file)
                        else:
                            documento.add_paragraph("[ERROR AL GENERAR IMAGEN DESDE HTML]")
                    else:
                        texto_limpio = limpiar_respuesta_final(respuesta_ia_bruta)
                        texto_corregido = corregir_numeracion_markdown(texto_limpio)
                        if texto_corregido: agregar_markdown_a_word(documento, texto_corregido)
                else:
                    generation_successful = True

        except Exception as e:
            st.error(f"Ocurrió un error crítico durante la generación del cuerpo: {e}"); generation_successful = False

        if generation_successful:
            project_name = st.session_state.selected_project['name']
            safe_project_name = re.sub(r'[\\/*?:"<>|]', "", project_name).replace(' ', '_')
            
            lot_name_clean_filename = lot_name_clean.replace(' ', '_')
            nombre_archivo_final = f"Cuerpo_Memoria_{safe_project_name}_{lot_name_clean_filename}.docx"
            
            doc_io = io.BytesIO()
            documento.save(doc_io)
            doc_io.seek(0)
            
            st.session_state.generated_doc_buffer = doc_io
            st.session_state.generated_doc_filename = nombre_archivo_final
            st.success("¡Cuerpo del documento generado con éxito!")
            st.rerun()

    if st.session_state.get("generated_doc_buffer"):
        st.info("El cuerpo del documento está listo para descargar o para el ensamblaje final.")
        st.download_button(
            label="📄 Descargar Cuerpo del Documento (.docx)",
            data=st.session_state.generated_doc_buffer,
            file_name=st.session_state.generated_doc_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    
    st.markdown("---")
    col_nav1, col_nav2 = st.columns(2)
    with col_nav1: 
        st.button("← Volver a Fase 4 (Plan de Prompts)", on_click=go_to_phase4, use_container_width=True)
    with col_nav2: 
        st.button(
            "Ir a Ensamblaje Final (F6) →", 
            on_click=go_to_phase6, 
            use_container_width=True, 
            type="primary", 
            disabled=not st.session_state.get("generated_doc_buffer")
        )
        
# =============================================================================
#           PÁGINA FASE 6: ENSAMBLAJE FINAL
# =============================================================================
def phase_6_page(model, go_to_phase5, back_to_project_selection_and_cleanup):
    st.markdown("<h3>FASE 6: Ensamblaje del Documento Final</h3>", unsafe_allow_html=True)
    
    selected_lot_text = "Análisis General"
    if st.session_state.get('selected_lot') and st.session_state.selected_lot != OPCION_ANALISIS_GENERAL:
        selected_lot_text = st.session_state.selected_lot
    
    st.info(f"Ensamblando la memoria técnica para: **{selected_lot_text}**")
    st.markdown("Este es el último paso. Se añadirá un índice y una introducción profesional al documento.")
    st.markdown("---")

    if not st.session_state.get("generated_doc_buffer"):
        st.warning("No se ha encontrado un documento de la Fase 5. Por favor, completa la fase anterior.")
        if st.button("← Ir a Fase 5"): 
            go_to_phase5()
            st.rerun()
        return

    if not st.session_state.get("generated_structure"):
        st.warning("No se ha encontrado la estructura del proyecto. Vuelve a una fase anterior para generarla.")
        return

    if st.button("🚀 Ensamblar Documento Final con Índice e Introducción", type="primary", use_container_width=True):
        try:
            with st.spinner("Ensamblando la versión definitiva..."):
                buffer_fase5 = st.session_state.generated_doc_buffer
                buffer_fase5.seek(0)
                documento_fase5 = docx.Document(buffer_fase5)
                
                texto_completo_original = "\n".join([p.text for p in documento_fase5.paragraphs if p.text.strip()])
                
                st.toast("Generando introducción estratégica...")
                idioma_seleccionado = st.session_state.get('project_language', 'Español')
                prompt_intro_formateado = PROMPT_GENERAR_INTRODUCCION.format(idioma=idioma_seleccionado)
                response_intro = model.generate_content([prompt_intro_formateado, texto_completo_original])
                introduccion_markdown = limpiar_respuesta_final(response_intro.text)
                
                st.toast("Creando documento final...")
                documento_final = docx.Document()
                estructura_memoria = st.session_state.generated_structure.get('estructura_memoria', [])
                
                generar_indice_word(documento_final, estructura_memoria)
                documento_final.add_page_break()
                documento_final.add_heading("Introducción", level=1)
                agregar_markdown_a_word(documento_final, corregir_numeracion_markdown(introduccion_markdown))
                documento_final.add_page_break()
                
                for element in documento_fase5.element.body:
                    documento_final.element.body.append(element)
                    
                doc_io_final = io.BytesIO()
                documento_final.save(doc_io_final)
                doc_io_final.seek(0)
                
                st.session_state.refined_doc_buffer = doc_io_final
                original_filename = st.session_state.generated_doc_filename
                st.session_state.refined_doc_filename = original_filename.replace("Cuerpo_", "Version_Final_")
                
                st.toast("Guardando versión final en Google Drive...")
                service = st.session_state.drive_service
                project_folder_id = st.session_state.selected_project['id']
                active_lot_folder_id = get_or_create_lot_folder_id(service, project_folder_id, lot_name=st.session_state.get('selected_lot'))
                
                doc_io_final.name = st.session_state.refined_doc_filename
                doc_io_final.type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                upload_file_to_drive(service, doc_io_final, active_lot_folder_id)

                st.success("¡Documento final ensamblado y guardado en Drive!")
                st.rerun()

        except Exception as e:
            st.error(f"Ocurrió un error crítico durante el ensamblaje final: {e}")

    if st.session_state.get("refined_doc_buffer"):
        st.balloons()
        st.success("¡Tu memoria técnica definitiva está lista!")
        st.download_button(
            label="🏆 Descargar Versión Definitiva (.docx)",
            data=st.session_state.refined_doc_buffer,
            file_name=st.session_state.refined_doc_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    st.markdown("---")
    col_nav1, col_nav2 = st.columns(2)
    with col_nav1:
        st.button("← Volver a Fase 5 (Redacción)", on_click=go_to_phase5, use_container_width=True)
    with col_nav2:
        st.button("✅ PROCESO FINALIZADO (Volver a selección de proyecto)", on_click=back_to_project_selection_and_cleanup, use_container_width=True, type="primary")



































